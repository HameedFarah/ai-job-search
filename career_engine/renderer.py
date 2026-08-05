from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .ats_styles import DEFAULT_STYLE_ID, available_style_ids, resolve_style, style_list
from .bundle import load_bundle
from .config import load_config
from .template import status as approved_template_status


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _libreoffice_binary(home: Path | None = None) -> str:
    """Resolve LibreOffice without requiring a system-wide installation or PATH change.

    Resolution order: explicit ``CAREER_ENGINE_LIBREOFFICE`` env override, then any
    system ``libreoffice``/``soffice`` on PATH, then the newest user-local
    ``~/.local/opt/libreoffice-*/opt/libreoffice*/program/soffice``. The user-local
    glob avoids mutating the global PATH and works for non-root installs such as
    ``~/.local/opt/libreoffice-26.2.5``.
    """
    candidates: list[str] = []
    override = os.environ.get("CAREER_ENGINE_LIBREOFFICE", "").strip()
    if override:
        candidates.append(override)
    for name in ("libreoffice", "soffice"):
        resolved = shutil.which(name)
        if resolved:
            candidates.append(resolved)
            continue
        # ``shutil.which`` relies on ``os.access(..., X_OK)`` and therefore may
        # reject an executable file located on a ``noexec`` mount. Preserve PATH
        # precedence by checking permission bits directly before falling back to
        # the user-local LibreOffice installation.
        for directory in os.environ.get("PATH", "").split(os.pathsep):
            if not directory:
                continue
            path_candidate = Path(directory) / name
            if path_candidate.is_file() and path_candidate.stat().st_mode & 0o111:
                candidates.append(str(path_candidate))
                break
    user_local = sorted(
        (home or Path.home()).glob(".local/opt/libreoffice-*/opt/libreoffice*/program/soffice"),
        reverse=True,
    )
    candidates.extend(str(path) for path in user_local)
    for candidate in candidates:
        path = Path(candidate).expanduser()
        # Check the executable permission bits directly. ``os.access(..., X_OK)``
        # also reflects mount-level ``noexec`` policy, which caused valid bundled
        # binaries and isolated test fixtures to be rejected even though the file
        # itself was correctly marked executable.
        if path.is_file() and path.stat().st_mode & 0o111:
            return str(path.resolve())
    return ""


def _libreoffice_profile_dir(home: Path | None = None) -> Path:
    """Return a writable directory for isolated headless LibreOffice user profiles.

    A fresh profile per conversion avoids lock collisions between concurrent headless
    instances. The profile and LibreOffice's own scratch space must not depend on a
    small or full system temp filesystem (e.g. a 100% full ``/tmp`` tmpfs), which
    otherwise aborts conversion with a write error; the user cache directory lives on
    the main disk and is writable for non-root installs.
    """
    base = (home or Path.home()) / ".cache" / "career-engine" / "libreoffice-profiles"
    base.mkdir(parents=True, exist_ok=True)
    return base


def render_tooling() -> dict[str, Any]:
    """Report the external tooling required for DOCX -> PDF rendering and PDF verification."""
    libreoffice = _libreoffice_binary()
    pdfinfo = shutil.which("pdfinfo")
    pdftotext = shutil.which("pdftotext")
    return {
        "libreoffice": bool(libreoffice),
        "libreoffice_path": libreoffice,
        "pdfinfo": bool(pdfinfo),
        "pdftotext": bool(pdftotext),
        "pdf_conversion_available": bool(libreoffice),
        "pdf_verification_available": bool(pdfinfo and pdftotext),
    }


def template_status(*, root: Path | None = None) -> dict[str, Any]:
    return approved_template_status(root=root)


def _all_paragraphs(document: Document) -> Iterable[Paragraph]:
    # Deduplicate on the lxml element itself, never on id() values: lxml proxy
    # objects are garbage-collected eagerly and Python reuses id() slots, which
    # caused placeholder paragraphs (metric boxes, evidence cards, achievements,
    # headline) to be skipped and the renderer to fail on the approved template.
    seen: set[Any] = set()

    def walk_table(table: Any) -> Iterable[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph._p not in seen:
                        seen.add(paragraph._p)
                        yield paragraph
                for nested in cell.tables:
                    yield from walk_table(nested)

    for paragraph in document.paragraphs:
        if paragraph._p not in seen:
            seen.add(paragraph._p)
            yield paragraph
    for table in document.tables:
        yield from walk_table(table)


def _replace_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _insert_paragraph_after(source: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(source._p)
    source._p.addnext(new_p)
    paragraph = Paragraph(new_p, source._parent)
    _replace_text(paragraph, text)
    return paragraph


def _claim_map(packet: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {claim["id"]: claim for claim in packet.get("selected_claims", [])}


def _metric_label(claim: dict[str, Any]) -> str:
    value = str(claim.get("value", "")).strip()
    label = str(claim.get("label", "")).strip()
    if value and label.lower().startswith(value.lower()):
        label = label[len(value):].strip(" -|:")
    return label or str(claim.get("safe_wording", ""))


def _recognized_context(packet: dict[str, Any]) -> str:
    text = " ".join(str(claim.get("safe_wording", "")) for claim in packet.get("selected_claims", []))
    names = [name for name in ("Saudi Aramco", "TotalEnergies", "Amazon") if name.lower() in text.lower()]
    sectors: list[str] = []
    low = text.lower()
    for token, label in (
        ("healthcare", "healthcare"),
        ("entertainment", "entertainment"),
        ("sports", "sports infrastructure"),
        ("government", "confidential government programmes"),
        ("logistics", "logistics and fulfilment"),
    ):
        if token in low and label not in sectors:
            sectors.append(label)
    values = names + sectors
    if not values:
        values = ["major Saudi design, supervision and project-delivery assignments"]
    if len(values) == 1:
        return values[0]
    return ", ".join(values[:-1]) + " and " + values[-1]


def build_render_input(job_id: str, generated_application: dict[str, Any], packet: dict[str, Any], *, root: Path | None = None) -> dict[str, Any]:
    config, paths = load_config(root)
    status = template_status(root=root)
    if not status["valid"]:
        raise FileNotFoundError("Approved Career Engine DOCX template is missing or invalid: " + status["path"])
    artifact_dir = paths.tracker_base / "artifacts" / job_id
    artifact_dir.mkdir(parents=True, exist_ok=True)
    render_input = {
        "schema_version": 1,
        "job_id": job_id,
        "bundle_hash": packet["bundle_hash"],
        "template": status,
        "layout": {
            "page_limit": config["template"]["page_limit"],
            "headshot_required": config["template"]["headshot_required"],
            "body_alignment": config["template"]["body_alignment"],
        },
        "outward_filename": packet["outward_filename"],
        "content": generated_application,
        "application_route": packet["application_route"],
    }
    path = artifact_dir / "render_input.json"
    path.write_text(json.dumps(render_input, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {"render_input": render_input, "path": str(path)}


def render_docx(job_id: str, generated_application: dict[str, Any], packet: dict[str, Any], *, root: Path | None = None) -> dict[str, Any]:
    config, paths = load_config(root)
    template = paths.repo_root / config["template"]["repository_path"]
    status = template_status(root=root)
    if not status["valid"]:
        raise FileNotFoundError("Approved Career Engine DOCX template is missing or invalid")
    artifact_dir = paths.tracker_base / "artifacts" / job_id
    artifact_dir.mkdir(parents=True, exist_ok=True)
    outward_pdf = packet["outward_filename"]
    outward_docx = Path(outward_pdf).with_suffix(".docx").name
    destination = artifact_dir / outward_docx
    document = Document(template)
    paragraphs = list(_all_paragraphs(document))

    headline = str(generated_application["headline"]).strip()
    profile = str(generated_application["leadership_profile"]["text"]).strip()
    current_bullets = [str(item["text"]).strip() for item in generated_application["current_role_bullets"]]
    if len(current_bullets) != 7:
        raise ValueError(f"The approved template requires exactly seven generated current-role bullets, got {len(current_bullets)}")
    claim_map = _claim_map(packet)
    metric_ids = generated_application["metric_claim_ids"]
    metrics = [claim_map[item] for item in metric_ids]

    achievement_paragraphs = [p for p in paragraphs if p.text.startswith("• [ACHIEVEMENT")]
    if len(achievement_paragraphs) != 7:
        raise ValueError(f"Template achievement placeholders changed: expected 7, got {len(achievement_paragraphs)}")
    for paragraph, text in zip(achievement_paragraphs, current_bullets):
        _replace_text(paragraph, "• " + text.lstrip("• ").strip())
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.08

    # The approved template originally carried two fixed current-role statements.
    # Remove them so every role bullet in the outward CV comes from the validated,
    # claim-cited generated application.
    for paragraph in list(paragraphs):
        if paragraph.text.startswith("• Directed end-to-end design and project delivery") or paragraph.text.startswith(
            "• Led multidisciplinary teams, consultants, contractors and stakeholders"
        ):
            _remove_paragraph(paragraph)

    earlier_items = generated_application.get("earlier_role_bullets", [])
    if len(earlier_items) != 11:
        raise ValueError(f"The approved template requires exactly eleven generated earlier-role bullets, got {len(earlier_items)}")

    def bullet_for_claim(claim_id: str) -> str:
        matches = [
            str(item["text"]).strip()
            for item in earlier_items
            if claim_id in item.get("claim_ids", [])
        ]
        if len(matches) != 1:
            raise ValueError(f"Expected exactly one earlier-role bullet citing {claim_id}, got {len(matches)}")
        return matches[0]

    cube_claim_order = [
        "cube.projects.25",
        "cube.team.10plus",
        "cube.assets.office_scale",
        "cube.agreements.30plus",
        "cube.bim.workflow.50",
        "cube.value_engineering.15plus",
        "cube.procurement.tender",
    ]
    cube_placeholder_starts = [
        "• Led multidisciplinary design delivery across 25 developments",
        "• Directed a core team of 10+ designers and engineers",
        "• Delivered major commercial assets",
        "• Negotiated and formalised 30+",
        "• Led BIM/Revit adoption",
        "• Delivered construction-cost savings above 15%",
        "• Managed procurement and tender strategy",
    ]
    cube_placeholders = [
        next((p for p in paragraphs if p.text.startswith(start)), None)
        for start in cube_placeholder_starts
    ]
    if any(paragraph is None for paragraph in cube_placeholders):
        raise ValueError("Template Cube Architects chronology placeholders changed")
    for paragraph, claim_id in zip(cube_placeholders, cube_claim_order):
        assert paragraph is not None
        _replace_text(paragraph, "• " + bullet_for_claim(claim_id).lstrip("• ").strip())
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(14)
        paragraph.paragraph_format.line_spacing = 1.08

    earlier_role_map = [
        ("Coordinated architectural design, construction documentation", "earlier.cube_project_architect.delivery"),
        ("Managed a 750-sqm retail branch", "earlier.cud.smartbuy.750"),
        ("Prepared feasibility studies and budgets", "earlier.procurement.20plus"),
        ("Developed coordinated design, tender and construction packages", "earlier.sigma.design_packages"),
    ]
    for start, claim_id in earlier_role_map:
        paragraph = next((p for p in paragraphs if p.text.startswith(start)), None)
        if paragraph is None:
            raise ValueError(f"Template earlier-career placeholder changed: {start}")
        _replace_text(paragraph, bullet_for_claim(claim_id))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.08

    replacements = {
        "[TARGET ROLE HEADLINE]": headline,
        "[TAILORED PROFILE: 55 TO 75 WORDS. POSITION THE CANDIDATE FOR THE VACANCY, USE THE OFFICIAL TITLE ONLY IN CHRONOLOGY, AND LEAD WITH THE STRONGEST EVIDENCE-SUPPORTED VALUE PROPOSITION.]": profile,
    }
    for paragraph in paragraphs:
        if paragraph.text in replacements:
            _replace_text(paragraph, replacements[paragraph.text])
            if paragraph.text == profile:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.08

    metric_value_paragraphs = [p for p in paragraphs if re.fullmatch(r"\[M[1-6]\]", p.text)]
    metric_label_paragraphs = [p for p in paragraphs if p.text.startswith("[VACANCY-RELEVANT")]
    if len(metric_value_paragraphs) != 6 or len(metric_label_paragraphs) != 6:
        raise ValueError("Template metric placeholders changed")
    for paragraph, claim in zip(metric_value_paragraphs, metrics):
        _replace_text(paragraph, str(claim.get("value", "")))
    for paragraph, claim in zip(metric_label_paragraphs, metrics):
        _replace_text(paragraph, _metric_label(claim))

    metric_set = set(metric_ids)
    evidence_claims = [
        claim for claim in packet.get("selected_claims", [])
        if claim.get("id") not in metric_set
        and not str(claim.get("id", "")).startswith(("credential.", "education.", "cube.", "earlier."))
    ]
    if len(evidence_claims) < 4:
        evidence_claims.extend(claim for claim in metrics if claim not in evidence_claims)
    evidence_claims = evidence_claims[:4]
    evidence_titles = [p for p in paragraphs if p.text.startswith("[EVIDENCE CARD")]
    evidence_texts = [p for p in paragraphs if p.text.startswith("[One concise evidence statement")]
    if len(evidence_titles) != 4 or len(evidence_texts) != 4:
        raise ValueError("Template evidence-card placeholders changed")
    for index, (title_p, text_p, claim) in enumerate(zip(evidence_titles, evidence_texts, evidence_claims), start=1):
        _replace_text(title_p, str(claim.get("label", f"Selected evidence {index}")))
        title_p.paragraph_format.space_after = Pt(3)
        _replace_text(text_p, str(claim.get("safe_wording", "")))
        text_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        text_p.paragraph_format.space_after = Pt(8)
        text_p.paragraph_format.line_spacing = 1.08

    for paragraph in paragraphs:
        if paragraph.text.startswith("Representative KSA context:"):
            _replace_text(paragraph, "Representative KSA context: " + _recognized_context(packet) + ".")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            break

    page_two_headings = {
        "EARLIER CAREER",
        "EDUCATION",
        "PROFESSIONAL CREDENTIALS & AFFILIATIONS",
        "LANGUAGES & NATIONALITY",
    }
    page_two_institutions = {
        "Zigurat Institute / Barcelona University, Spain",
        "New York Institute of Technology, USA | GPA 3.93/4.00, with honors",
        "University of Jordan, Jordan | GPA 3.14/4.00",
    }
    for paragraph in paragraphs:
        if paragraph.text in page_two_headings:
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
        elif paragraph.text in page_two_institutions:
            paragraph.paragraph_format.space_after = Pt(10)

    document.core_properties.title = f"Abdelhamid Farah - {headline}"
    document.core_properties.subject = "Vacancy-tailored curriculum vitae"
    document.core_properties.author = "Abdelhamid Farah"
    document.core_properties.keywords = "architecture, design management, project delivery, Saudi Arabia"
    document.save(destination)
    return {
        "docx": str(destination),
        "sha256": file_sha256(destination),
        "template_sha256": status["sha256"],
        "outward_filename": outward_docx,
    }


def copy_template_to_workspace(job_id: str, *, root: Path | None = None) -> dict[str, Any]:
    config, paths = load_config(root)
    template = paths.repo_root / config["template"]["repository_path"]
    if not template.is_file():
        raise FileNotFoundError(template)
    destination = paths.tracker_base / "artifacts" / job_id / template.name
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, destination)
    return {"path": str(destination), "sha256": file_sha256(destination)}


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> dict[str, Any]:
    libreoffice = _libreoffice_binary()
    if not libreoffice:
        return {"converted": False, "blocker": "libreoffice_missing"}
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="career-engine-lo-", dir=_libreoffice_profile_dir()) as profile_dir:
        profile = Path(profile_dir).resolve()
        # Keep LibreOffice's own scratch space on the main disk next to the fresh
        # user profile: a full system /tmp would otherwise abort the export with a
        # write error even though the output directory itself has space.
        env = dict(os.environ)
        env["TMPDIR"] = str(profile)
        completed = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--norestore",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
            env=env,
        )
    pdf = output_dir / (docx_path.stem + ".pdf")
    return {
        "converted": completed.returncode == 0 and pdf.is_file(),
        "returncode": completed.returncode,
        "libreoffice": libreoffice,
        "pdf": str(pdf) if pdf.is_file() else "",
        "stdout": completed.stdout[-2000:],
        "stderr": completed.stderr[-2000:],
        "sha256": file_sha256(pdf) if pdf.is_file() else "",
    }


def _command_output(command: list[str], *, timeout: int = 120) -> tuple[int, str, str]:
    completed = subprocess.run(command, capture_output=True, text=True, timeout=timeout, check=False)
    return completed.returncode, completed.stdout, completed.stderr


def verify_pdf(pdf_path: Path, *, root: Path | None = None) -> dict[str, Any]:
    config, _ = load_config(root)
    bundle = load_bundle(root)
    findings: list[dict[str, str]] = []
    page_count = 0
    if shutil.which("pdfinfo"):
        code, out, err = _command_output(["pdfinfo", str(pdf_path)])
        if code == 0:
            match = re.search(r"^Pages:\s+(\d+)", out, re.MULTILINE)
            if match:
                page_count = int(match.group(1))
        else:
            findings.append({"severity": "error", "code": "pdfinfo_failed", "message": err[-500:]})
    if page_count != int(config["template"]["page_limit"]):
        findings.append({"severity": "error", "code": "page_count", "message": f"Expected 2 pages, got {page_count}"})
    text = ""
    if shutil.which("pdftotext"):
        code, out, err = _command_output(["pdftotext", "-layout", str(pdf_path), "-"])
        if code == 0:
            text = out
        else:
            findings.append({"severity": "error", "code": "pdftotext_failed", "message": err[-500:]})
    if not text.strip():
        findings.append({"severity": "error", "code": "missing_text_layer", "message": "PDF text layer is empty"})
    for value in bundle["config"]["policy"].get("prohibited_experience_names", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "prohibited_name", "message": value})
    for value in bundle["config"]["policy"].get("prohibited_terms", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "prohibited_term", "message": value})
    for value in bundle["config"]["policy"].get("availability_patterns", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "availability", "message": value})
    for value in bundle["config"]["policy"].get("forbidden_characters", []):
        if value in text:
            findings.append({"severity": "error", "code": "forbidden_character", "message": value})
    required = ["Abdelhamid Farah", "hameedfarah@gmail.com", "+966 53 079 6449", "Consultant"]
    for value in required:
        if value.lower() not in text.lower():
            findings.append({"severity": "error", "code": "required_text_missing", "message": value})
    return {
        "valid": not any(item["severity"] == "error" for item in findings),
        "pdf": str(pdf_path),
        "sha256": file_sha256(pdf_path),
        "page_count": page_count,
        "text_characters": len(text),
        "findings": findings,
    }


def render_and_verify(job_id: str, generated_application: dict[str, Any], packet: dict[str, Any], *, root: Path | None = None) -> dict[str, Any]:
    docx = render_docx(job_id, generated_application, packet, root=root)
    docx_path = Path(docx["docx"])
    converted = convert_docx_to_pdf(docx_path, docx_path.parent)
    if not converted.get("converted"):
        return {"valid": False, "docx": docx, "conversion": converted, "verification": {}}
    verification = verify_pdf(Path(converted["pdf"]), root=root)
    return {"valid": verification["valid"], "docx": docx, "conversion": converted, "verification": verification}


# ---------------------------------------------------------------------------
# ATS Linear renderer
#
# The ATS Linear template is a single-column, ATS-optimized linear resume:
# no sidebar, no headshot, no tables, no text boxes, no floating objects and
# no images. It is built programmatically in reusable code (guaranteeing the
# ATS-safe layout by construction) from the template spec in
# ``projects/job-automation/config/ats-linear-template.v1.json``, which is a
# NON-bundle source: it is deliberately not part of the compiled runtime bundle
# inputs so existing generation packets keep their bundle_hash valid. All
# content comes from the validated generated_application.json and
# generation_packet.json (claim-cited content) plus verified facts already
# carried by the approved Modern Executive Sidebar template and the canonical
# Vault profile - nothing is invented.
# ---------------------------------------------------------------------------

_ATS_MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def ats_template_status(*, root: Path | None = None) -> dict[str, Any]:
    """Load and validate the ATS Linear template spec (non-bundle config)."""
    _, paths = load_config(root)
    path = paths.ats_template_path
    if not path.is_file():
        return {"valid": False, "present": False, "path": str(path), "reason": "missing"}
    try:
        spec = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        return {"valid": False, "present": True, "path": str(path), "reason": f"invalid_json: {exc}"}
    required = {"schema_version", "template_id", "version", "page_limit", "fonts", "margins_cm", "section_headings"}
    missing = sorted(required - set(spec))
    if missing:
        return {"valid": False, "present": True, "path": str(path), "reason": f"missing_fields: {missing}"}
    return {
        "valid": True,
        "present": True,
        "path": str(path),
        "spec": spec,
        "template_id": spec["template_id"],
        "version": spec["version"],
    }


def _ats_month_year(value: str) -> str:
    value = str(value).strip()
    if not value or value.lower() == "present":
        return "Present"
    parts = value.split("-")
    if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
        try:
            return f"{_ATS_MONTHS[int(parts[1]) - 1]} {parts[0]}"
        except IndexError:
            pass
    return value


def _ats_outward_filename(packet_outward: str) -> str:
    """Turn the executive outward filename into the ATS-named variant.

    ``Abdelhamid_Farah_CV_Senior_Design_Manager.pdf`` becomes
    ``Abdelhamid_Farah_CV_Senior_Design_Manager_ATS.pdf``.
    """
    stem, suffix = os.path.splitext(packet_outward)
    return f"{stem}_ATS{suffix}"


def _ats_experience_blocks(application: dict[str, Any], packet: dict[str, Any]) -> list[dict[str, Any]]:
    """Order the claim-cited bullets into the verified six-role chronology."""
    chronology = packet.get("career_chronology") or []
    if len(chronology) != 6:
        raise ValueError(f"ATS Linear expects exactly six verified career-chronology roles, got {len(chronology)}")
    current_bullets = [str(item["text"]).strip() for item in application["current_role_bullets"]]
    if len(current_bullets) != 7:
        raise ValueError(f"ATS Linear requires exactly seven generated current-role bullets, got {len(current_bullets)}")
    earlier = application.get("earlier_role_bullets", [])
    if len(earlier) != 11:
        raise ValueError(f"ATS Linear requires exactly eleven generated earlier-role bullets, got {len(earlier)}")
    earlier_texts = [str(item["text"]).strip() for item in earlier]
    expected_single_claims = [
        "earlier.cube_project_architect.delivery",
        "earlier.cud.smartbuy.750",
        "earlier.procurement.20plus",
        "earlier.sigma.design_packages",
    ]
    for index, expected in enumerate(expected_single_claims, start=7):
        if expected not in earlier[index].get("claim_ids", []):
            raise ValueError(f"ATS Linear earlier-role bullet {index} does not cite {expected}")
    blocks: list[dict[str, Any]] = []
    cube_bullets = earlier_texts[0:7]
    for position, role in enumerate(chronology):
        if position == 0:
            bullets = current_bullets
        elif position == 1:
            bullets = cube_bullets
        else:
            bullets = earlier_texts[position + 5:position + 6]
        blocks.append({
            "title": str(role.get("title", "")).strip(),
            "employer": str(role.get("employer", "")).strip(),
            "location": str(role.get("location", "")).strip(),
            "start": str(role.get("start", "")).strip(),
            "end": str(role.get("end", "")).strip(),
            "bullets": bullets,
        })
    return blocks


def _ats_language_line(languages: list[str]) -> str:
    parts: list[str] = []
    for item in languages:
        if " - " in item:
            lang, level = item.split(" - ", 1)
            parts.append(f"{lang.strip()}: {level.strip().capitalize()}")
        else:
            parts.append(str(item))
    return " | ".join(parts)


def _ats_build_document(spec: dict[str, Any], application: dict[str, Any], packet: dict[str, Any],
                        style: dict[str, Any] | None = None) -> Document:
    """Build the ATS Linear document.

    ``style`` is a resolved design-option spec from ``ats_styles.resolve_style``.
    When ``None`` (or omitted) the document is built exactly as the existing
    conservative ATS Linear visual (``ats-classic``): same fonts, margins,
    spacing, rule color and section order, so the default output is unchanged.
    """
    if style is None:
        style = resolve_style(DEFAULT_STYLE_ID, spec)
    fonts = style["fonts"]
    font_name = str(fonts["name"])
    margins = style["margins_cm"]
    colors = style["colors"]
    line_spacing = float(style["line_spacing"])
    section_order = style["section_order"]
    presentation = style["presentation"]

    def _hex_color(hex_value: str) -> RGBColor | None:
        if not hex_value or hex_value == "000000":
            return None
        try:
            return RGBColor.from_string(hex_value)
        except ValueError:
            return None

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(fonts["body_size_pt"])
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attribute), font_name)

    def styled_run(paragraph: Paragraph, text: str, *, size: float, bold: bool = False,
                   color: str | None = None) -> None:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = font_name
        run.font.size = Pt(size)
        hex_color = _hex_color(color or "000000")
        if hex_color is not None:
            run.font.color.rgb = hex_color
        run_rpr = run._element.get_or_add_rPr()
        run_rfonts = run_rpr.get_or_add_rFonts()
        for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            run_rfonts.set(qn(attribute), font_name)

    def add_paragraph(text: str = "", *, size: float, bold: bool = False, align: int | None = None,
                      space_before: float = 0, space_after: float = 0, left_indent: float = 0,
                      first_line_indent: float = 0, keep_with_next: bool = False,
                      color: str | None = None) -> Paragraph:
        paragraph = document.add_paragraph()
        if text:
            styled_run(paragraph, text, size=size, bold=bold, color=color)
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.keep_with_next = keep_with_next
        if left_indent:
            paragraph.paragraph_format.left_indent = Cm(left_indent)
        if first_line_indent:
            paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
        return paragraph

    def _add_bottom_rule(paragraph: Paragraph, *, size: int, space: int, color: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), str(space))
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def add_section_heading(text: str, *, size_delta: float = 0.0) -> None:
        size = float(fonts["section_heading_size_pt"]) + size_delta
        rule = presentation.get("section_rule", {"enabled": True})
        paragraph = add_paragraph(
            text, size=size, bold=True,
            space_before=presentation.get("section_before", 8),
            space_after=presentation.get("section_after", 3),
            keep_with_next=True,
            color=colors["section_heading"],
        )
        if rule.get("enabled"):
            _add_bottom_rule(
                paragraph,
                size=int(rule.get("size", 6)),
                space=int(rule.get("space", 2)),
                color=colors["section_rule"],
            )

    def add_bullet(text: str, *, size: float | None = None, space_after: float | None = None,
                   color: str | None = None) -> None:
        add_paragraph(
            "• " + text,
            size=size if size is not None else fonts["bullet_size_pt"],
            space_after=space_after if space_after is not None else presentation.get("bullet_after", 2),
            left_indent=presentation.get("bullet_left_indent", 0.5),
            first_line_indent=presentation.get("bullet_hanging_indent", -0.5),
            color=color,
        )

    def add_role(block: dict[str, Any]) -> None:
        meta = " | ".join(part for part in (
            block["employer"],
            block["location"],
            f"{_ats_month_year(block['start'])} - {_ats_month_year(block['end'])}",
        ) if part)
        title_paragraph = add_paragraph(
            block["title"], size=fonts["role_title_size_pt"], bold=True,
            space_before=presentation.get("role_before", 5),
            space_after=presentation.get("role_after", 0),
            keep_with_next=True,
            color=colors["role_title"],
        )
        if presentation.get("role_rule"):
            _add_bottom_rule(
                title_paragraph,
                size=int(presentation.get("role_rule_size", 4)),
                space=int(presentation.get("role_rule_space", 2)),
                color=colors["section_rule"],
            )
        if meta:
            add_paragraph(meta, size=fonts["role_meta_size_pt"],
                          space_after=presentation.get("role_meta_after", 2),
                          keep_with_next=True, color=colors["role_meta"])
        for bullet in block["bullets"]:
            add_bullet(bullet)

    identity = packet.get("identity") or {}
    name_alignment = WD_ALIGN_PARAGRAPH.CENTER if presentation.get("name_alignment") == "center" else WD_ALIGN_PARAGRAPH.LEFT

    def render_header() -> None:
        add_paragraph(str(identity.get("professional_name") or "Abdelhamid Farah"),
                      size=fonts["name_size_pt"], bold=True,
                      align=name_alignment, space_after=presentation.get("name_after", 1),
                      color=colors["name"])
        contact_parts = [
            str(identity.get("location", "")).strip(),
            str(identity.get("ksa_phone", "")).strip(),
            str(identity.get("outward_email", "")).strip(),
            "linkedin.com/in/abd-farah",
        ]
        add_paragraph(" | ".join(part for part in contact_parts if part),
                      size=fonts["contact_size_pt"],
                      align=name_alignment, space_after=presentation.get("contact_after", 4))
        add_paragraph(str(application["headline"]).strip(), size=fonts["headline_size_pt"], bold=True,
                      align=name_alignment, space_after=presentation.get("headline_after", 2),
                      color=colors["headline"])

    def render_summary() -> None:
        add_section_heading("PROFESSIONAL SUMMARY")
        profile = str(application["leadership_profile"]["text"]).strip()
        add_paragraph(profile, size=fonts["body_size_pt"], space_after=presentation.get("profile_after", 2))

    def render_experience() -> None:
        add_section_heading("PROFESSIONAL EXPERIENCE", size_delta=presentation.get("experience_heading_size_delta_pt", 0.0))
        for block in _ats_experience_blocks(application, packet):
            add_role(block)

    def render_education() -> None:
        add_section_heading("EDUCATION")
        for entry in spec["verified_sections"]["education"]:
            degree = str(entry["degree"])
            institution = str(entry["institution"])
            year = str(entry["year"])
            add_paragraph(degree, size=fonts["body_size_pt"], bold=True,
                          space_after=presentation.get("education_degree_after", 0),
                          keep_with_next=True)
            add_paragraph(f"{institution} | {year}", size=fonts["body_size_pt"],
                          space_after=presentation.get("education_institution_after", 2))

    def render_credentials() -> None:
        add_section_heading("PROFESSIONAL CREDENTIALS")
        for credential in spec["verified_sections"]["credentials"]:
            add_bullet(str(credential))

    def render_languages() -> None:
        add_section_heading("LANGUAGES")
        languages = [str(item) for item in (identity.get("languages") or [])]
        add_paragraph(_ats_language_line(languages), size=fonts["body_size_pt"],
                      space_after=presentation.get("languages_after", 1))
        nationalities = [str(item) for item in (identity.get("nationalities") or [])]
        if nationalities:
            add_paragraph("Nationality: " + " and ".join(nationalities), size=fonts["body_size_pt"],
                          space_after=presentation.get("nationality_after", 2))

    sections = {
        "summary": render_summary,
        "experience": render_experience,
        "education": render_education,
        "credentials": render_credentials,
        "languages": render_languages,
    }

    render_header()
    for key in section_order:
        sections[key]()

    return document


def render_ats_docx(job_id: str, generated_application: dict[str, Any], packet: dict[str, Any], *,
                    root: Path | None = None, style_id: str = DEFAULT_STYLE_ID,
                    out_dir: Path | str | None = None, filename: str | None = None) -> dict[str, Any]:
    config, paths = load_config(root)
    status = ats_template_status(root=root)
    if not status["valid"]:
        raise FileNotFoundError("ATS Linear template spec is missing or invalid: " + status.get("reason", status.get("path", "")))
    spec = status["spec"]
    style = resolve_style(style_id, spec)
    if out_dir is None:
        artifact_dir = paths.tracker_base / "artifacts" / job_id
    else:
        artifact_dir = Path(out_dir)
    artifact_dir.mkdir(parents=True, exist_ok=True)
    outward = _ats_outward_filename(packet["outward_filename"])
    docx_name = Path(filename).name if filename else Path(outward).with_suffix(".docx").name
    docx_path = artifact_dir / docx_name
    document = _ats_build_document(spec, generated_application, packet, style=style)
    document.core_properties.title = f"Abdelhamid Farah - {generated_application['headline']}"
    document.core_properties.subject = "Vacancy-tailored curriculum vitae - ATS linear"
    document.core_properties.author = "Abdelhamid Farah"
    document.core_properties.keywords = "architecture, design management, project delivery, Saudi Arabia, ATS"
    document.save(str(docx_path))
    return {
        "docx": str(docx_path),
        "sha256": file_sha256(docx_path),
        "template_id": status["template_id"],
        "template_version": status["version"],
        "style_id": style_id,
        "outward_filename": docx_path.name,
    }


def ats_docx_checks(docx_path: Path) -> dict[str, Any]:
    """Structural ATS-safety checks on a generated ATS Linear DOCX."""
    findings: list[dict[str, str]] = []
    document = Document(docx_path)
    if document.tables:
        findings.append({"severity": "error", "code": "tables_present", "message": f"ATS Linear DOCX contains {len(document.tables)} tables"})
    if document.inline_shapes:
        findings.append({"severity": "error", "code": "images_present", "message": f"ATS Linear DOCX contains {len(document.inline_shapes)} inline images/shapes"})
    xml = document.element.xml
    if "<w:txbxContent" in xml:
        findings.append({"severity": "error", "code": "text_boxes_present", "message": "ATS Linear DOCX contains text boxes"})
    if "<wp:anchor" in xml:
        findings.append({"severity": "error", "code": "floating_objects_present", "message": "ATS Linear DOCX contains floating objects"})
    if "<w:drawing" in xml or "<w:pict" in xml:
        findings.append({"severity": "error", "code": "drawing_present", "message": "ATS Linear DOCX contains drawing/image content"})
    sections = document.sections
    if len(sections) != 1:
        findings.append({"severity": "error", "code": "multiple_sections", "message": f"Expected one section, got {len(sections)}"})
    for index, section in enumerate(sections):
        cols = section._sectPr.find(qn("w:cols"))
        num = int(cols.get(qn("w:num"))) if cols is not None and cols.get(qn("w:num")) else 1
        if num != 1:
            findings.append({"severity": "error", "code": "multi_column", "message": f"Section {index} uses {num} columns"})
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if not text.strip():
        findings.append({"severity": "error", "code": "empty_document", "message": "DOCX contains no selectable text"})
    return {
        "valid": not any(item["severity"] == "error" for item in findings),
        "findings": findings,
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "paragraph_count": len(document.paragraphs),
        "text_characters": len(text),
    }


def _pdf_word_boxes(pdf_path: Path) -> dict[int, list[tuple[float, float, float, float]]]:
    """Return per-page word bounding boxes (xMin, yMin, xMax, yMax) via pdftotext -bbox."""
    pages: dict[int, list[tuple[float, float, float, float]]] = {}
    if not shutil.which("pdftotext"):
        return pages
    code, out, _err = _command_output(["pdftotext", "-bbox", str(pdf_path), "-"], timeout=120)
    if code != 0 or not out.strip():
        return pages
    page_re = re.compile(r'<page width="[\d.]+" height="[\d.]+">(.*?)</page>', re.DOTALL)
    word_re = re.compile(r'<word xMin="([\d.]+)" yMin="([\d.]+)" xMax="([\d.]+)" yMax="([\d.]+)"', re.DOTALL)
    for index, match in enumerate(page_re.finditer(out), start=1):
        pages[index] = [
            (float(a), float(b), float(c), float(d))
            for a, b, c, d in word_re.findall(match.group(1))
        ]
    return pages


def _pdf_column_findings(pdf_path: Path, *, gap_pt: float = 80.0) -> list[dict[str, str]]:
    """Detect multi-column text layouts from word x-coordinates.

    A single-column resume clusters every word start near the left margin
    (body plus hanging-indent bullets). A second text column creates a second
    x0 cluster separated by a gap wider than ``gap_pt``.
    """
    findings: list[dict[str, str]] = []
    pages = _pdf_word_boxes(pdf_path)
    for page_index, words in sorted(pages.items()):
        if len(words) < 20:
            continue
        clusters: list[list[float]] = []
        for value in sorted(word[0] for word in words):
            if clusters and value - clusters[-1][-1] <= gap_pt:
                clusters[-1].append(value)
            else:
                clusters.append([value])
        significant = [cluster for cluster in clusters if len(cluster) >= 4]
        if len(significant) > 1:
            findings.append({
                "severity": "error",
                "code": "multi_column_layout",
                "message": f"Page {page_index} shows {len(significant)} text-column clusters",
            })
    return findings


def verify_ats_pdf(pdf_path: Path, *, root: Path | None = None, docx_path: Path | None = None) -> dict[str, Any]:
    """ATS-safe verification of a rendered ATS Linear PDF (+ its DOCX when given).

    Page count is reported as a warning when it differs from the 2-page target:
    ATS integrity and complete relevant chronology take precedence over forcing
    the page count. Every other check (text layer, policy terms, images,
    columns, DOCX structure) is an error when violated.
    """
    config, _ = load_config(root)
    bundle = load_bundle(root)
    status = ats_template_status(root=root)
    spec = status.get("spec", {})
    page_limit = int(spec.get("page_limit", config["template"]["page_limit"]))
    findings: list[dict[str, str]] = []
    page_count = 0
    if shutil.which("pdfinfo"):
        code, out, err = _command_output(["pdfinfo", str(pdf_path)])
        if code == 0:
            match = re.search(r"^Pages:\s+(\d+)", out, re.MULTILINE)
            if match:
                page_count = int(match.group(1))
        else:
            findings.append({"severity": "error", "code": "pdfinfo_failed", "message": err[-500:]})
    if page_count != page_limit:
        findings.append({
            "severity": "warning",
            "code": "page_count",
            "message": f"Expected {page_limit} pages, got {page_count}; ATS integrity takes precedence over forcing the page count",
        })
    text = ""
    if shutil.which("pdftotext"):
        code, out, err = _command_output(["pdftotext", "-layout", str(pdf_path), "-"])
        if code == 0:
            text = out
        else:
            findings.append({"severity": "error", "code": "pdftotext_failed", "message": err[-500:]})
    if not text.strip():
        findings.append({"severity": "error", "code": "missing_text_layer", "message": "PDF text layer is empty"})
    for value in bundle["config"]["policy"].get("prohibited_experience_names", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "prohibited_name", "message": value})
    for value in bundle["config"]["policy"].get("prohibited_terms", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "prohibited_term", "message": value})
    for value in bundle["config"]["policy"].get("availability_patterns", []):
        if value.lower() in text.lower():
            findings.append({"severity": "error", "code": "availability", "message": value})
    for value in bundle["config"]["policy"].get("forbidden_characters", []):
        if value in text:
            findings.append({"severity": "error", "code": "forbidden_character", "message": value})
    required = ["Abdelhamid Farah", "hameedfarah@gmail.com", "+966 53 079 6449", "Consultant"]
    for value in required:
        if value.lower() not in text.lower():
            findings.append({"severity": "error", "code": "required_text_missing", "message": value})
    if shutil.which("pdfimages"):
        code, out, _err = _command_output(["pdfimages", "-list", str(pdf_path)])
        image_lines = [
            line for line in out.splitlines()
            if line.strip() and not line.startswith(("page", "----"))
        ]
        if image_lines:
            findings.append({
                "severity": "error",
                "code": "pdf_images_present",
                "message": f"PDF embeds {len(image_lines)} raster images",
            })
    findings.extend(_pdf_column_findings(pdf_path))
    docx_result: dict[str, Any] = {}
    if docx_path is not None and Path(docx_path).is_file():
        docx_result = ats_docx_checks(Path(docx_path))
        findings.extend(docx_result["findings"])
    return {
        "valid": not any(item["severity"] == "error" for item in findings),
        "pdf": str(pdf_path),
        "sha256": file_sha256(pdf_path),
        "page_count": page_count,
        "word_count": len(text.split()),
        "text_characters": len(text),
        "findings": findings,
        "docx_checks": docx_result,
    }


def render_ats_and_verify(job_id: str, generated_application: dict[str, Any], packet: dict[str, Any], *,
                          root: Path | None = None, style_id: str = DEFAULT_STYLE_ID) -> dict[str, Any]:
    docx = render_ats_docx(job_id, generated_application, packet, root=root, style_id=style_id)
    docx_path = Path(docx["docx"])
    converted = convert_docx_to_pdf(docx_path, docx_path.parent)
    if not converted.get("converted"):
        return {"valid": False, "template": "ats-linear", "style_id": style_id, "docx": docx, "conversion": converted, "verification": {}}
    verification = verify_ats_pdf(Path(converted["pdf"]), root=root, docx_path=docx_path)
    return {"valid": verification["valid"], "template": "ats-linear", "style_id": style_id, "docx": docx, "conversion": converted, "verification": verification}


# ---------------------------------------------------------------------------
# ATS design options gallery
#
# Renders all five reusable ATS-safe visual design options from the *same*
# verified generated_application.json and generation_packet.json content into
# ``projects/job-automation/artifacts/ats-design-options/`` with DOCX, PDF,
# first-page PNG preview and a machine-readable manifest.json. Existing portal
# filenames and the ats-linear recommendation rule are never touched.
# ---------------------------------------------------------------------------


def ats_style_list() -> dict[str, Any]:
    """Registry of the reusable ATS design style options."""
    entries = style_list()
    return {
        "default_style_id": DEFAULT_STYLE_ID,
        "count": len(entries),
        "styles": entries,
        "portal_recommendation": "ats-linear (unchanged)",
    }


def _pdf_raster_safety_findings(pdf_path: Path) -> dict[str, Any]:
    """Detect word-box overlap and out-of-bounds clipping from pdftotext -bbox.

    This is a raster-adjacent safety pass used by the design-options gallery:
    word bounding boxes that intersect (suggesting text collision) or extend
    beyond the page (suggesting clipping) are reported as warnings. The strict
    ATS gate remains ``verify_ats_pdf``.
    """
    result: dict[str, Any] = {"overlap": [], "clipping": [], "clean": True}
    if not shutil.which("pdftotext"):
        return result
    code, out, _err = _command_output(["pdftotext", "-bbox", str(pdf_path), "-"], timeout=120)
    if code != 0 or not out.strip():
        return result
    page_re = re.compile(r'<page width="([\d.]+)" height="([\d.]+)">(.*?)</page>', re.DOTALL)
    word_re = re.compile(r'<word xMin="([\d.]+)" yMin="([\d.]+)" xMax="([\d.]+)" yMax="([\d.]+)"', re.DOTALL)
    for page_index, match in enumerate(page_re.finditer(out), start=1):
        width = float(match.group(1))
        height = float(match.group(2))
        words = [
            (float(a), float(b), float(c), float(d))
            for a, b, c, d in word_re.findall(match.group(3))
        ]
        for a, b, c, d in words:
            if a < -0.5 or b < -0.5 or c > width + 0.5 or d > height + 0.5:
                result["clipping"].append({
                    "page": page_index,
                    "box": [round(a, 2), round(b, 2), round(c, 2), round(d, 2)],
                    "page_size_pt": [round(width, 2), round(height, 2)],
                })
        for i in range(len(words)):
            ax1, ay1, ax2, ay2 = words[i]
            for j in range(i + 1, len(words)):
                bx1, by1, bx2, by2 = words[j]
                overlap_x = min(ax2, bx2) - max(ax1, bx1)
                overlap_y = min(ay2, by2) - max(ay1, by1)
                if overlap_x > 0 and overlap_y > 0 and overlap_x * overlap_y > 1.0:
                    result["overlap"].append({
                        "page": page_index,
                        "box_a": [round(ax1, 2), round(ay1, 2), round(ax2, 2), round(ay2, 2)],
                        "box_b": [round(bx1, 2), round(by1, 2), round(bx2, 2), round(by2, 2)],
                    })
    result["clean"] = not result["overlap"] and not result["clipping"]
    return result


def _png_page_preview(pdf_path: Path, out_dir: Path, stem: str, *, dpi: int = 110) -> dict[str, Any]:
    """Render the first PDF page to PNG with pdftoppm."""
    if not shutil.which("pdftoppm"):
        return {"preview": "", "rendered": False, "blocker": "pdftoppm_missing"}
    out_dir.mkdir(parents=True, exist_ok=True)
    temporary = out_dir / (stem + ".preview")
    code, _out, err = _command_output(
        ["pdftoppm", "-f", "1", "-l", "1", "-singlefile", "-png", "-r", str(dpi), str(pdf_path), str(temporary)],
        timeout=120,
    )
    produced = out_dir / (stem + ".preview.png")
    if code != 0 or not produced.is_file():
        return {"preview": "", "rendered": False, "blocker": err[-500:] or "pdftoppm_failed"}
    target = out_dir / f"{stem}.page1.png"
    if target.exists():
        target.unlink()
    produced.replace(target)
    return {"preview": target.name, "rendered": True, "path": str(target), "dpi": dpi}


def _png_pixel_raster_checks(png_path: Path, margins_cm: dict[str, Any], *, dpi: int = 110) -> dict[str, Any]:
    """Pixel-level raster inspection of a preview PNG.

    Verifies (a) no ink in the outermost page band (content clipped at the page
    edge), (b) the configured side margins are respected (ink-free margin
    zones) and (c) the page is not blank. Word-level overlap and out-of-bounds
    clipping are checked separately against all pages via ``pdftotext -bbox``
    in ``_pdf_raster_safety_findings``.
    """
    try:
        from PIL import Image
    except ImportError:
        return {"checked": False, "blocker": "PIL_missing", "findings": []}
    image = Image.open(str(png_path)).convert("L")
    width, height = image.size
    pixels = image.load()

    def _cm_to_px(cm: float) -> int:
        return int(round(float(cm) / 2.54 * dpi))

    def _band_ink(start_x: int, end_x: int, start_y: int, end_y: int) -> int:
        ink = 0
        for y in range(max(0, start_y), min(height, end_y)):
            for x in range(max(0, start_x), min(width, end_x)):
                if pixels[x, y] < 200:
                    ink += 1
        return ink

    findings: list[dict[str, Any]] = []
    edge_band = _cm_to_px(0.35)
    for name, band in (
        ("left_edge", (0, edge_band, 0, height)),
        ("right_edge", (width - edge_band, width, 0, height)),
        ("top_edge", (0, width, 0, edge_band)),
        ("bottom_edge", (0, width, height - edge_band, height)),
    ):
        ink = _band_ink(*band)
        if ink > 0:
            findings.append({"severity": "warning", "code": "edge_ink", "edge": name, "ink_pixels": ink})

    tolerance_px = max(_cm_to_px(0.3), 1)
    for name, margin_cm in (("left", margins_cm.get("left")), ("right", margins_cm.get("right")),
                            ("top", margins_cm.get("top")), ("bottom", margins_cm.get("bottom"))):
        if margin_cm is None:
            continue
        zone = max(_cm_to_px(margin_cm) - tolerance_px, 1)
        if name == "left":
            band = (0, zone, 0, height)
        elif name == "right":
            band = (width - zone, width, 0, height)
        elif name == "top":
            band = (0, width, 0, zone)
        else:
            band = (0, width, height - zone, height)
        ink = _band_ink(*band)
        if ink > 0:
            findings.append({
                "severity": "warning", "code": "margin_ink",
                "margin": name, "margin_cm": margin_cm, "ink_pixels": ink,
            })

    total_ink = _band_ink(0, width, 0, height)
    return {
        "checked": True,
        "width": width,
        "height": height,
        "dpi": dpi,
        "total_ink_pixels": total_ink,
        "blank": total_ink == 0,
        "findings": findings,
        "clean": not findings and total_ink > 0,
    }


def render_ats_design_options(job_id: str, *, root: Path | None = None,
                              out_dir: Path | str | None = None,
                              include_previews: bool = True) -> dict[str, Any]:
    """Render, verify and preview all five ATS design options for a job package.

    Reads the validated ``generated_application.json`` and
    ``generation_packet.json`` for ``job_id`` (the content used by the portal
    pipeline) and produces one DOCX + PDF + page-1 PNG per style plus a
    ``manifest.json`` gallery. Portal artifacts are never modified.
    """
    config, paths = load_config(root)
    artifact_dir = paths.tracker_base / "artifacts" / job_id
    application_path = artifact_dir / "generated_application.json"
    packet_path = artifact_dir / "generation_packet.json"
    if not application_path.is_file() or not packet_path.is_file():
        return {"job_id": job_id, "valid": False, "blocker": "generated_application_missing"}
    application = json.loads(application_path.read_text(encoding="utf-8"))
    packet = json.loads(packet_path.read_text(encoding="utf-8"))

    if out_dir is None:
        out_dir = paths.tracker_base / "artifacts" / "ats-design-options"
    out_dir = Path(out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    status = ats_template_status(root=root)
    spec = status.get("spec", {}) if status.get("valid") else {}
    stem = Path(_ats_outward_filename(packet["outward_filename"])).stem

    classic_paragraphs: list[str] | None = None
    entries: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    for style_id in available_style_ids():
        style = resolve_style(style_id, spec)
        slug = style_id.replace("ats-", "")
        base = f"{stem}_{slug}"

        docx_result = render_ats_docx(
            job_id, application, packet, root=root,
            style_id=style_id, out_dir=out_dir, filename=f"{base}.docx",
        )
        docx_path = Path(docx_result["docx"])
        docx_texts = [re.sub(r"\s+", " ", p.text).strip() for p in Document(str(docx_path)).paragraphs]
        content: dict[str, Any] = {"paragraph_count": len(docx_texts)}
        if classic_paragraphs is None:
            classic_paragraphs = docx_texts
            content["matches_classic"] = True
        else:
            content["matches_classic"] = sorted(classic_paragraphs) == sorted(docx_texts)
        content["word_count_docx"] = len(" ".join(text for text in docx_texts if text).split())

        converted = convert_docx_to_pdf(docx_path, out_dir)
        if not converted.get("converted"):
            errors.append({
                "style_id": style_id,
                "blocker": "pdf_conversion_failed",
                "stderr": str(converted.get("stderr", ""))[-500:],
            })
            entries.append({
                "id": style_id,
                "label": style["label"],
                "description": style["description"],
                "recommended_for_portal": style["recommended_for_portal"],
                "docx": docx_path.name,
                "pdf": "",
                "preview": "",
                "ats": {"valid": False, "findings": [{"severity": "error", "code": "pdf_conversion_failed"}]},
                "content_integrity": content,
            })
            continue

        pdf_path = Path(converted["pdf"])
        verification = verify_ats_pdf(pdf_path, root=root, docx_path=docx_path)
        raster = _pdf_raster_safety_findings(pdf_path)
        preview = _png_page_preview(pdf_path, out_dir, base) if include_previews else {"preview": "", "rendered": False}
        if preview.get("rendered"):
            pixel = _png_pixel_raster_checks(out_dir / preview["preview"], style["margins_cm"], dpi=int(preview.get("dpi", 110)))
        else:
            pixel = {"checked": False, "clean": False, "findings": [], "blocker": preview.get("blocker", "preview_not_rendered")}
        raster_clean = bool(raster["clean"]) and bool(pixel.get("clean"))
        entries.append({
            "id": style_id,
            "label": style["label"],
            "description": style["description"],
            "recommended_for_portal": style["recommended_for_portal"],
            "docx": docx_path.name,
            "pdf": pdf_path.name,
            "preview": preview.get("preview", ""),
            "ats": {
                "valid": verification["valid"],
                "findings": verification["findings"],
                "docx_checks_valid": bool(verification.get("docx_checks", {}).get("valid", True)),
                "tables": int(verification.get("docx_checks", {}).get("tables", 0)),
                "images": int(verification.get("docx_checks", {}).get("images", 0)),
                "page_count": int(verification.get("page_count", 0)),
                "word_count": int(verification.get("word_count", 0)),
                "text_characters": int(verification.get("text_characters", 0)),
                "sha256": verification.get("sha256", ""),
            },
            "raster": {
                "clean": raster["clean"],
                "overlap_findings": raster["overlap"],
                "clipping_findings": raster["clipping"],
            },
            "content_integrity": content,
            "visual_inspection": {
                "method": "programmatic raster inspection: pdftotext -bbox (all pages) + PIL pixel checks (page-1 PNG)",
                "pdf_bbox": {"clean": bool(raster["clean"]), "overlap_count": len(raster["overlap"]), "clipping_count": len(raster["clipping"])},
                "pixel": pixel,
                "clipping_or_overlap": "none_detected" if raster_clean else "findings_present",
            },
        })

    all_ats_safe = bool(entries) and all(entry.get("ats", {}).get("valid") for entry in entries)
    page_counts = [int(entry.get("ats", {}).get("page_count", 0)) for entry in entries]
    all_two_pages = bool(page_counts) and all(count == 2 for count in page_counts)
    try:
        directory = str(out_dir.relative_to(paths.repo_root))
    except ValueError:
        directory = str(out_dir)
    manifest = {
        "schema_version": 1,
        "kind": "ats-design-options-gallery",
        "title": "Five ATS Resume Design Options",
        "sample_job_id": job_id,
        "sample_headline": str(application.get("headline", "")),
        "bundle_hash": str(packet.get("bundle_hash", "")),
        "default_style_id": DEFAULT_STYLE_ID,
        "recommended_style": DEFAULT_STYLE_ID,
        "portal_recommendation": "ats-linear (unchanged)",
        "directory": directory,
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "styles": entries,
        "summary": {
            "style_count": len(entries),
            "all_ats_safe": all_ats_safe,
            "all_two_pages": all_two_pages,
            "variants_needing_three_pages": [entry["id"] for entry in entries if int(entry.get("ats", {}).get("page_count", 0)) == 3],
            "content_identical_to_classic": all(entry.get("content_integrity", {}).get("matches_classic", False) for entry in entries),
            "errors": errors,
        },
        "valid": all_ats_safe and not errors,
    }
    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest
