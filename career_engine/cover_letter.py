from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .config import load_config
from .renderer import convert_docx_to_pdf, file_sha256


def _filename_token(value: str) -> str:
    token = re.sub(r"[^A-Za-z0-9]+", "_", str(value or "").strip()).strip("_")
    return token or "Application"


def _verify_cover_letter_pdf(pdf: Path) -> dict[str, Any]:
    findings: list[dict[str, str]] = []
    page_count = 0
    text = ""
    if shutil.which("pdfinfo"):
        completed = subprocess.run(["pdfinfo", str(pdf)], capture_output=True, text=True, timeout=60, check=False)
        if completed.returncode == 0:
            match = re.search(r"^Pages:\s+(\d+)", completed.stdout, re.MULTILINE)
            if match:
                page_count = int(match.group(1))
        else:
            findings.append({"severity": "error", "code": "pdfinfo_failed", "message": completed.stderr[-500:]})
    if page_count not in {1, 2}:
        findings.append({"severity": "error", "code": "page_count", "message": f"Expected a 1-2 page cover letter, got {page_count}"})
    if shutil.which("pdftotext"):
        completed = subprocess.run(["pdftotext", "-layout", str(pdf), "-"], capture_output=True, text=True, timeout=60, check=False)
        if completed.returncode == 0:
            text = completed.stdout
        else:
            findings.append({"severity": "error", "code": "pdftotext_failed", "message": completed.stderr[-500:]})
    if not text.strip():
        findings.append({"severity": "error", "code": "missing_text_layer", "message": "Cover-letter PDF text layer is empty"})
    return {
        "valid": not any(item["severity"] == "error" for item in findings),
        "pdf": str(pdf),
        "sha256": file_sha256(pdf),
        "page_count": page_count,
        "text_characters": len(text),
        "findings": findings,
    }


def render_cover_letter(
    job_id: str,
    generated_application: dict[str, Any],
    packet: dict[str, Any],
    *,
    root: Path | None = None,
) -> dict[str, Any]:
    """Render validated cover prose into a standalone application letter.

    The generated-application schema already requires ``cover_email`` with
    claim-cited, validated prose for every route. Reusing that canonical prose
    avoids a second uncontrolled writing pass while ensuring portal/ATS packages
    also carry a real cover-letter artifact.
    """
    _, paths = load_config(root)
    artifact_dir = paths.tracker_base / "artifacts" / job_id
    artifact_dir.mkdir(parents=True, exist_ok=True)

    vacancy = packet.get("vacancy") or {}
    identity = packet.get("identity") or {}
    role = str(vacancy.get("role", "Position")).strip()
    company = str(vacancy.get("company", "")).strip()
    filename = f"Abdelhamid_Farah_Cover_Letter_{_filename_token(role)}.docx"
    destination = artifact_dir / filename

    email = generated_application.get("cover_email") or {}
    body = str(email.get("body", "")).strip()
    if not body:
        raise ValueError("Validated generated application has no cover_email.body")

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run(str(identity.get("professional_name") or identity.get("name") or "Abdelhamid Farah"))
    run.bold = True
    run.font.size = Pt(16)

    contacts = [
        str(identity.get("location", "")).strip(),
        str(identity.get("ksa_phone", "")).strip(),
        str(identity.get("outward_email", "")).strip(),
        "linkedin.com/in/abd-farah",
    ]
    contact = document.add_paragraph(" | ".join(item for item in contacts if item))
    contact.paragraph_format.space_after = Pt(18)

    if company:
        company_p = document.add_paragraph(company)
        company_p.paragraph_format.space_after = Pt(4)
    subject = document.add_paragraph(f"Re: Application for {role}")
    subject.runs[0].bold = True
    subject.paragraph_format.space_after = Pt(14)

    for block in re.split(r"\n\s*\n", body):
        text = block.strip()
        if not text:
            continue
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(9)
        paragraph.paragraph_format.line_spacing = 1.08

    document.core_properties.title = f"Abdelhamid Farah - Cover Letter - {role}"
    document.core_properties.subject = f"Application for {role}{' at ' + company if company else ''}"
    document.core_properties.author = "Abdelhamid Farah"
    document.save(destination)
    return {
        "docx": str(destination),
        "sha256": file_sha256(destination),
        "outward_filename": destination.name,
    }


def render_cover_letter_and_verify(
    job_id: str,
    generated_application: dict[str, Any],
    packet: dict[str, Any],
    *,
    root: Path | None = None,
) -> dict[str, Any]:
    docx = render_cover_letter(job_id, generated_application, packet, root=root)
    docx_path = Path(docx["docx"])
    conversion = convert_docx_to_pdf(docx_path, docx_path.parent)
    if not conversion.get("converted"):
        return {"valid": False, "docx": docx, "conversion": conversion, "verification": {}}
    verification = _verify_cover_letter_pdf(Path(conversion["pdf"]))
    return {"valid": verification["valid"], "docx": docx, "conversion": conversion, "verification": verification}
