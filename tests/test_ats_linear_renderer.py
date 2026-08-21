from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from career_engine.bundle import build_bundle
from career_engine.core import decide_route, match_evidence, normalize_job, score_fit
from career_engine.generation import create_generation_packet
from career_engine.renderer import (
    _ats_language_line,
    _ats_month_year,
    _ats_outward_filename,
    ats_docx_checks,
    ats_template_status,
    render_ats_and_verify,
    render_ats_docx,
    render_tooling,
    verify_ats_pdf,
)
from tests.test_career_engine_v1 import engine_root, job_payload, valid_application  # noqa: F401

REPO = Path(__file__).resolve().parents[1]


_CHRONOLOGY = [
    {"employer": "Arab Sustainable Architecture / Tubaila Team Workshop (TTW)", "location": "Saudi Arabia", "title": "District Manager", "start": "2022-12", "end": "present"},
    {"employer": "Cube Architects", "location": "Jordan", "title": "Design Manager / Senior Architect", "start": "2013-03", "end": "2022-11"},
    {"employer": "Cube Architects", "location": "Jordan", "title": "Project Architect", "start": "2012-01", "end": "2013-02"},
    {"employer": "Creative Urban Designs", "location": "Jordan", "title": "Project Manager", "start": "2010-09", "end": "2011-12"},
    {"employer": "Al-Mehanya Real Estate Investment", "location": "Jordan", "title": "Project Architect", "start": "2008-06", "end": "2010-08"},
    {"employer": "Sigma Consulting Engineers", "location": "Jordan", "title": "Architect", "start": "2004-06", "end": "2008-02"},
]


def _packet(job_payload: dict, engine_root: Path) -> dict:
    bundle = build_bundle(engine_root)
    normalized = normalize_job(job_payload, bundle["taxonomy"])
    matches = match_evidence(normalized, bundle)
    score = score_fit(normalized, matches, bundle)
    route = decide_route(normalized, bundle)
    packet = create_generation_packet(job_id="ats-render-test-job", normalized_job=normalized, matches=matches, score=score, route=route, bundle=bundle)
    # Mirror the production packet: full six-role verified chronology and identity
    packet["career_chronology"] = [dict(entry) for entry in _CHRONOLOGY]
    packet["identity"] = {
        "professional_name": "Abdelhamid Farah",
        "location": "Riyadh, Saudi Arabia",
        "nationalities": ["Jordanian", "Brazilian"],
        "languages": ["Arabic - native", "English - fluent"],
        "outward_email": "hameedfarah@gmail.com",
        "ksa_phone": "+966 53 079 6449",
        "current_role": "District Manager",
        "current_employer": "Arab Sustainable Architecture / Tubaila Team Workshop (TTW)",
    }
    return packet


def _application(packet: dict) -> dict:
    return valid_application(packet)


def test_ats_template_spec_loads_from_non_bundle_config(engine_root: Path) -> None:
    status = ats_template_status(root=engine_root)
    assert status["valid"] is True
    assert status["present"] is True
    assert status["template_id"] == "ats-linear"
    spec = status["spec"]
    assert spec["headshot_required"] is False
    assert spec["ats_safe_guarantees"]["one_column"] is True
    assert spec["ats_safe_guarantees"]["no_tables"] is True
    assert spec["ats_safe_guarantees"]["no_images"] is True
    assert "PROFESSIONAL EXPERIENCE" in spec["section_headings"]
    assert spec["route_recommendation"]["portal"] == "ats-linear"
    assert spec["route_recommendation"]["email"] == "modern-executive-sidebar"


def test_ats_spec_is_not_a_runtime_bundle_source(engine_root: Path) -> None:
    """Adding the ATS spec must not invalidate the compiled bundle hash."""
    bundle = build_bundle(engine_root)
    config_path = engine_root / "projects/job-automation/config/career-engine.v1.json"
    config = json.loads(config_path.read_text(encoding="utf-8"))
    assert "ats-linear-template.v1.json" not in json.dumps(config)
    status = build_bundle(engine_root)
    assert status["bundle_hash"] == bundle["bundle_hash"]


def test_ats_outward_filename_includes_ats_suffix() -> None:
    assert _ats_outward_filename("Abdelhamid_Farah_CV_Senior_Design_Manager.pdf") == "Abdelhamid_Farah_CV_Senior_Design_Manager_ATS.pdf"
    assert _ats_outward_filename("Abdelhamid_Farah_CV_Design_Manager.pdf").endswith("_ATS.pdf")
    assert _ats_month_year("2022-12") == "Dec 2022"
    assert _ats_month_year("present") == "Present"
    assert _ats_language_line(["Arabic - native", "English - fluent"]) == "Arabic: Native | English: Fluent"


def test_render_ats_docx_creates_linear_document(job_payload: dict, engine_root: Path) -> None:
    packet = _packet(job_payload, engine_root)
    application = _application(packet)
    result = render_ats_docx("ats-render-test-job", application, packet, root=engine_root)
    path = Path(result["docx"])
    assert path.is_file()
    assert result["template_id"] == "ats-linear"
    assert result["outward_filename"] == "Abdelhamid_Farah_CV_Senior_Design_Governance_Manager_ATS.docx"

    document = Document(path)
    texts = [p.text for p in document.paragraphs]
    rendered = "\n".join(texts)
    assert application["headline"] in rendered
    assert application["leadership_profile"]["text"] in rendered
    assert "PROFESSIONAL SUMMARY" in texts
    assert "PROFESSIONAL EXPERIENCE" in texts
    assert "EDUCATION" in texts
    assert "PROFESSIONAL CREDENTIALS" in texts
    assert "LANGUAGES" in texts
    # Full verified chronology: every employer and title line is present
    assert "District Manager" in rendered
    assert "Cube Architects" in rendered
    assert "Creative Urban Designs" in rendered
    assert "Al-Mehanya Real Estate Investment" in rendered
    assert "Sigma Consulting Engineers" in rendered
    # All 7 current-role bullets and all 11 earlier-role bullets rendered
    for bullet in application["current_role_bullets"]:
        assert bullet["text"] in rendered
    for bullet in application["earlier_role_bullets"]:
        assert bullet["text"] in rendered
    # No template placeholders or leftover executive-only markers
    assert "[ACHIEVEMENT" not in rendered
    assert "[M1]" not in rendered
    assert "[TARGET ROLE HEADLINE]" not in rendered
    # Verified education facts
    assert "Global Master's in Construction Project Management" in rendered
    assert "Bachelor of Science in Architectural Engineering" in rendered
    # Credentials
    assert "Saudi Council of Engineers" in rendered
    assert "Contracts Management Professional" in rendered
    assert "SCE Classification: Consultant" in rendered


def test_render_ats_docx_is_ats_safe(job_payload: dict, engine_root: Path) -> None:
    packet = _packet(job_payload, engine_root)
    result = render_ats_docx("ats-render-test-job", _application(packet), packet, root=engine_root)
    checks = ats_docx_checks(Path(result["docx"]))
    assert checks["valid"] is True
    assert checks["tables"] == 0
    assert checks["images"] == 0
    codes = {item["code"] for item in checks["findings"]}
    assert "tables_present" not in codes
    assert "images_present" not in codes
    assert "text_boxes_present" not in codes
    assert "floating_objects_present" not in codes
    assert "multi_column" not in codes


def test_verify_ats_pdf_column_detection(tmp_path: Path, engine_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """The ATS-safe column check must detect a second text column and accept a single one."""
    from career_engine.renderer import _pdf_column_findings, _pdf_word_boxes

    # Single column: every word starts near the left margin (body + hanging-indent bullets)
    single_column = {
        1: [(48.3 + 0.5 * (i % 3), 40 + i * 12, 120 + i, 52 + i * 12) for i in range(40)],
    }
    # Two columns: a right-hand column starts ~300pt from the left margin
    two_column = {
        1: [(48.3 + 0.5 * (i % 3), 40 + i * 12, 120 + i, 52 + i * 12) for i in range(40)]
        + [(300 + 0.5 * (i % 3), 40 + i * 12, 420 + i, 52 + i * 12) for i in range(40)],
    }
    monkeypatch.setattr("career_engine.renderer._pdf_word_boxes", lambda pdf: single_column)
    assert _pdf_column_findings(tmp_path / "single.pdf") == []
    monkeypatch.setattr("career_engine.renderer._pdf_word_boxes", lambda pdf: two_column)
    findings = _pdf_column_findings(tmp_path / "two.pdf")
    assert any(item["code"] == "multi_column_layout" for item in findings)


@pytest.mark.skipif(
    not render_tooling()["pdf_conversion_available"]
    or not render_tooling()["pdf_verification_available"],
    reason="libreoffice or poppler tools unavailable",
)
def test_verify_ats_pdf_detects_embedded_image(tmp_path: Path, engine_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """ATS-safe PDF verification must flag embedded raster images."""
    from career_engine.renderer import _command_output
    from docx import Document as DocxDocument

    docx_path = tmp_path / "image.docx"
    document = DocxDocument()
    document.add_paragraph("Abdelhamid Farah")
    document.save(str(docx_path))

    class FakeTool:
        def __init__(self, command: list[str]) -> None:
            self.returncode = 0
            self.stderr = ""
            tool = command[0]
            if tool.endswith("pdfinfo"):
                self.stdout = "Pages:           2\nPage size: 595 x 842 pts (A4)"
            elif tool.endswith("pdftotext") and "-bbox" not in command:
                self.stdout = "Abdelhamid Farah Riyadh +966 53 079 6449 Consultant"
            elif tool.endswith("pdftotext") and "-bbox" in command:
                self.stdout = '<page width="595" height="842"><word xMin="48" yMin="50" xMax="100" yMax="65">Abdelhamid</word></page>'
            elif tool.endswith("pdfimages"):
                self.stdout = (
                    "page   num  type   width height color comp bpc  enc interp  object ID x-ppi y-ppi size ratio\n"
                    "--------------------------------------------------------------------------------------------\n"
                    "   1     0 image    400   300  rgb     3    8  jpeg   no        12  0    72    72  20K  2.3%\n"
                )
            else:
                self.stdout = ""

    monkeypatch.setattr(
        "career_engine.renderer._command_output",
        lambda command, timeout=120: (0, FakeTool(command).stdout, ""),
    )
    pdf_path = tmp_path / "image.pdf"
    pdf_path.write_bytes(b"fake pdf")
    result = verify_ats_pdf(pdf_path, root=engine_root, docx_path=docx_path)
    assert result["valid"] is False
    assert any(item["code"] == "pdf_images_present" for item in result["findings"])
    assert any(item["code"] == "required_text_missing" for item in result["findings"])


@pytest.mark.skipif(
    not render_tooling()["pdf_conversion_available"]
    or not render_tooling()["pdf_verification_available"],
    reason="libreoffice or poppler tools unavailable",
)
def test_render_ats_and_verify_full_pipeline(job_payload: dict, engine_root: Path) -> None:
    packet = _packet(job_payload, engine_root)
    application = _application(packet)
    result = render_ats_and_verify("ats-render-test-job", application, packet, root=engine_root)
    assert result["valid"] is True
    assert result["template"] == "ats-linear"
    verification = result["verification"]
    assert verification["page_count"] == 2
    assert verification["text_characters"] > 1000
    assert verification["word_count"] > 100
    error_codes = {item["code"] for item in verification["findings"] if item["severity"] == "error"}
    assert "missing_text_layer" not in error_codes
    assert "pdf_images_present" not in error_codes
    assert "multi_column_layout" not in error_codes
    assert "tables_present" not in error_codes
    # DOCX structural checks ran as part of verification
    assert verification["docx_checks"]["valid"] is True


def test_render_ats_requires_full_chronology(job_payload: dict, engine_root: Path) -> None:
    packet = _packet(job_payload, engine_root)
    application = _application(packet)
    packet = dict(packet)
    packet["career_chronology"] = packet["career_chronology"][:3]
    with pytest.raises(ValueError, match="exactly six"):
        render_ats_docx("ats-render-test-job", application, packet, root=engine_root)


def test_cli_render_ats_reports_missing_application(engine_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from career_engine.cli import main

    monkeypatch.setenv("CAREER_ENGINE_REPO_ROOT", str(engine_root))
    code = main(["render-ats", "--job-id", "missing-job-0000"])
    assert code == 10
