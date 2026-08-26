from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from tools.regenerate_sidebar_template import _norm_text, content_integrity_qa  # noqa: E402


def test_norm_text_strips_bidi_controls() -> None:
    value = "Abdelhamid\u200f Farah \u202a(مستشار)\u202c"
    assert _norm_text(value) == "abdelhamid farah (مستشار)"


def test_content_integrity_qa_tolerates_extraction_artifacts(tmp_path: Path) -> None:
    """pdftotext dehyphenation at line breaks and bidi reordering around RTL
    runs are rendering/extraction artifacts; QA must not flag them as lost
    content while genuinely missing text still fails."""
    from docx import Document

    expectations = {
        "headline": ["Design Governance Lead for giga-scale programmes"],
        "profile": [],
        "current_bullets": [],
        "metric_values": ["112+", "26+", "42%", "55.4%", "74", "30+"],
        "metric_labels": ["projects", "professionals", "turnaround", "win rate", "assignments", "agreements"],
        "evidence_titles": ["card one", "card two", "card three", "card four"],
        "evidence_texts": ["one", "two", "three", "four"],
        "ksa_context": [],
        "cube_bullets": [],
        "earlier_bullets": [],
        "identity": ["hameedfarah@gmail.com"],
    }
    document = Document()
    document.add_paragraph("Design Governance Lead for giga-")
    document.add_paragraph(
        "scale programmes, 112+ projects, hameedfarah@gmail.com "
        "26+ professionals 42% turnaround 55.4% win rate 74 assignments 30+ agreements"
    )
    for title, text in (("card one", "one"), ("card two", "two"), ("card three", "three"), ("card four", "four")):
        document.add_paragraph(title)
        document.add_paragraph(text)
    docx_path = tmp_path / "sample.docx"
    document.save(str(docx_path))
    # Simulated pdftotext output: hard-break dehyphenation ("gigascale") and
    # punctuation relocation around the RTL credential run.
    pdf_text = (
        "Design Governance Lead for gigascale programmes, (مستشار) 112+ "
        "projects, hameedfarah@gmail.com "
        "26+ professionals 42% turnaround 55.4% win rate 74 assignments 30+ agreements "
        "card one one card two two card three three card four four "
        "Cube Architects EARLIER CAREER EDUCATION PROFESSIONAL CREDENTIALS "
        "LANGUAGES Zigurat New York Institute of Technology University of Jordan SCE CMP"
    )
    bundle = {"config": {"policy": {"prohibited_experience_names": [], "prohibited_terms": []}}}
    import tools.regenerate_sidebar_template as tool

    original = tool._pdf_page_text
    tool._pdf_page_text = lambda _pdf: [pdf_text]
    try:
        report = content_integrity_qa(expectations, docx_path, tmp_path / "sample.pdf", bundle)
    finally:
        tool._pdf_page_text = original
    assert report["valid"] is True, report["findings"]

    # Genuinely missing content must still be flagged in both docx and pdf.
    expectations["metric_values"] = ["999+", "26+", "42%", "55.4%", "74", "30+"]
    tool._pdf_page_text = lambda _pdf: [pdf_text]
    try:
        report_missing = content_integrity_qa(expectations, docx_path, tmp_path / "sample.pdf", bundle)
    finally:
        tool._pdf_page_text = original
    codes = {finding["code"] for finding in report_missing["findings"]}
    assert codes == {"missing_in_docx", "missing_in_pdf"}


def test_content_integrity_qa_rejects_embedded_numeric_false_positives(tmp_path: Path) -> None:
    """Boundary hardening: a metric expectation must never be satisfied by an
    occurrence embedded inside unrelated numbers or words — 42% in 1420,
    55.4% in 1,554, '74 assignments' in 174, 112+ in 2112, one in components —
    while genuinely present values still pass."""
    from docx import Document

    expectations = {
        "headline": [],
        "profile": [],
        "current_bullets": [],
        "metric_values": ["42%", "55.4%", "74", "112+"],
        "metric_labels": ["turnaround", "win rate", "assignments", "projects"],
        "evidence_titles": ["one"],
        "evidence_texts": [],
        "ksa_context": [],
        "cube_bullets": [],
        "earlier_bullets": [],
        "identity": ["hameedfarah@gmail.com"],
    }
    traps = (
        "Delivery review: 1420 packages audited, 1,554 drawings checked, "
        "174 site visits, 2112 snags closed, key components inspected."
    )
    anchors = (
        " Cube Architects EARLIER CAREER EDUCATION PROFESSIONAL CREDENTIALS "
        "LANGUAGES Zigurat New York Institute of Technology University of Jordan SCE CMP"
    )
    document = Document()
    document.add_paragraph(traps)
    document.add_paragraph("hameedfarah@gmail.com")
    docx_path = tmp_path / "traps.docx"
    document.save(str(docx_path))
    bundle = {"config": {"policy": {"prohibited_experience_names": [], "prohibited_terms": []}}}
    import tools.regenerate_sidebar_template as tool

    original = tool._pdf_page_text
    tool._pdf_page_text = lambda _pdf: [f"{traps} hameedfarah@gmail.com{anchors}"]
    try:
        report = content_integrity_qa(expectations, docx_path, tmp_path / "traps.pdf", bundle)
    finally:
        tool._pdf_page_text = original
    assert report["valid"] is False
    missing = {finding["value"] for finding in report["findings"] if finding["code"].startswith("missing_in_")}
    assert {"42%", "55.4%", "74", "112+", "one"} <= missing


def test_content_integrity_qa_accepts_boundary_aligned_neighbors(tmp_path: Path) -> None:
    """The same metric values placed next to the trap-like numbers (genuine
    boundary-aligned occurrences) must still be accepted."""
    from docx import Document

    expectations = {
        "headline": [],
        "profile": [],
        "current_bullets": [],
        "metric_values": ["42%", "55.4%", "74", "112+", "26+", "30+"],
        "metric_labels": ["turnaround", "win rate", "assignments", "projects", "professionals", "agreements"],
        "evidence_titles": ["card one", "card two", "card three", "card four"],
        "evidence_texts": ["one", "two", "three", "four"],
        "ksa_context": [],
        "cube_bullets": [],
        "earlier_bullets": [],
        "identity": ["hameedfarah@gmail.com"],
    }
    body = (
        "Cost control delivered a 42% turnaround on fees with a 55.4% win "
        "rate across 74 assignments and 112+ projects with 26+ professionals "
        "and 30+ agreements; reference 1420 and 1,554 and 2112 remain "
        "unrelated contract numbers; card one one card two two card three "
        "three card four four. hameedfarah@gmail.com"
    )
    anchors = (
        " Cube Architects EARLIER CAREER EDUCATION PROFESSIONAL CREDENTIALS "
        "LANGUAGES Zigurat New York Institute of Technology University of Jordan SCE CMP"
    )
    document = Document()
    document.add_paragraph(body + anchors)
    docx_path = tmp_path / "aligned.docx"
    document.save(str(docx_path))
    bundle = {"config": {"policy": {"prohibited_experience_names": [], "prohibited_terms": []}}}
    import tools.regenerate_sidebar_template as tool

    original = tool._pdf_page_text
    tool._pdf_page_text = lambda _pdf: [body + anchors]
    try:
        report = content_integrity_qa(expectations, docx_path, tmp_path / "aligned.pdf", bundle)
    finally:
        tool._pdf_page_text = original
    assert report["valid"] is True, report["findings"]
