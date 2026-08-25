#!/usr/bin/env python3
"""Regenerate Modern Executive Sidebar revisions from the configured template.

Regeneration runs against the live CareerTracker/artifact authority: point
``CAREER_ENGINE_TRACKER_BASE`` at the authoritative ``projects/job-automation``
directory while running this tool from a clean source worktree. Every render
writes a NEW template-version-suffixed revision file and never overwrites an
existing artifact, so submitted evidence stays immutable. ATS Linear and cover
letter artifacts are never touched, and no email/send/submit path is invoked.

Subcommands:
  plan                  Enumerate jobs that carry a sidebar artifact.
  render --job-id X     Render the sidebar revision and run full QA.
  render --all [--limit N] [--only JOBID]  Batch render + QA.
  apply --job-id X      Record a passing QA'd revision in the tracker.
  apply --all           Record every job whose stored QA report passed.

QA evidence is written to ``<tracker_base>/runtime/sidebar-v15-qa/<job_id>.json``
(ignored runtime evidence, never committed to Git).
"""
from __future__ import annotations

import argparse
import json
import subprocess
import sys
import tempfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from career_engine.bundle import load_bundle  # noqa: E402
from career_engine.config import load_config  # noqa: E402
from career_engine.renderer import (  # noqa: E402
    _claim_map,
    _metric_label,
    _recognized_context,
    convert_docx_to_pdf,
    ensure_page_fit,
    file_sha256,
    render_docx,
    verify_pdf,
)

CUBE_CLAIM_ORDER = [
    "cube.projects.25",
    "cube.team.10plus",
    "cube.assets.office_scale",
    "cube.agreements.30plus",
    "cube.bim.workflow.50",
    "cube.value_engineering.15plus",
    "cube.procurement.tender",
]
EARLIER_CLAIM_MAP = [
    ("earlier.cube_project_architect.delivery",),
    ("earlier.cud.smartbuy.750",),
    ("earlier.procurement.20plus",),
    ("earlier.sigma.design_packages",),
]
FIXED_ANCHORS = [
    "Cube Architects",
    "EARLIER CAREER",
    "EDUCATION",
    "PROFESSIONAL CREDENTIALS",
    "LANGUAGES",
    "Zigurat",
    "New York Institute of Technology",
    "University of Jordan",
    "SCE",
    "CMP",
]
PLACEHOLDERS_ABSENT = [
    "[ACHIEVEMENT",
    "[M1]",
    "[M2]",
    "[M3]",
    "[M4]",
    "[M5]",
    "[M6]",
    "[VACANCY-RELEVANT",
    "[EVIDENCE CARD",
    "[One concise",
    "[TAILORED PROFILE",
    "[TARGET ROLE",
    "hameedo@gmail.com",
]
DASHES = ["\u2014", "\u2013"]  # em dash, en dash


def qa_dir(root: Path | None = None) -> Path:
    _, paths = load_config(root)
    return paths.tracker_base / "runtime" / "sidebar-v15-qa"


def _docx_text(docx_path: Path) -> str:
    from docx import Document

    document = Document(str(docx_path))
    texts: list[str] = []

    def walk_table(table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
                for inner in cell.tables:
                    walk_table(inner)

    texts.extend(p.text for p in document.paragraphs)
    for table in document.tables:
        walk_table(table)
    return "\n".join(texts)


def _pdf_page_text(pdf_path: Path) -> list[str]:
    # Reading-order extraction (no -layout): the sidebar column visually
    # interleaves with body lines in layout mode, which would split multi-line
    # body paragraphs with sidebar fragments and break contiguity checks.
    completed = subprocess.run(
        ["pdftotext", str(pdf_path), "-"],
        capture_output=True, text=True, timeout=120, check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError("pdftotext failed")
    return completed.stdout.split("\f")


def _norm(value: str) -> str:
    return " ".join(str(value).split()).casefold()


def _bullet_for_claim(claim_id: str, earlier_items: list[dict[str, Any]], claim_map: dict[str, dict[str, Any]], used: set[int]) -> str:
    for index, item in enumerate(earlier_items):
        if claim_id in item.get("claim_ids", []) and index not in used:
            used.add(index)
            return str(item["text"]).strip()
    claim = claim_map.get(claim_id) or {}
    safe = str(claim.get("safe_wording") or "").strip()
    if safe:
        return safe
    raise ValueError(f"No generated earlier-role bullet or verified safe wording is available for {claim_id}")


def source_expectations(application: dict[str, Any], packet: dict[str, Any], *, root: Path | None = None) -> dict[str, list[str]]:
    """Job-specific source strings that MUST survive into the rendered output."""
    claim_map = _claim_map(packet)
    metrics = [claim_map[item] for item in application["metric_claim_ids"]]
    metric_set = set(application["metric_claim_ids"])
    evidence_claims = [
        claim for claim in packet.get("selected_claims", [])
        if claim.get("id") not in metric_set
        and not str(claim.get("id", "")).startswith(("credential.", "education.", "cube.", "earlier."))
    ]
    if len(evidence_claims) < 4:
        evidence_claims.extend(claim for claim in metrics if claim not in evidence_claims)
    evidence_claims = evidence_claims[:4]

    earlier_items = application.get("earlier_role_bullets", [])
    used: set[int] = set()
    cube_bullets = [_bullet_for_claim(claim_id, earlier_items, claim_map, used) for claim_id in CUBE_CLAIM_ORDER]
    earlier_bullets = [_bullet_for_claim(entry[0], earlier_items, claim_map, used) for entry in EARLIER_CLAIM_MAP]

    identity = packet.get("identity") or {}
    required_identity = [
        str(identity.get("outward_email", "")).strip(),
        "+966 53 079 6449",
        "Abdelhamid Farah",
    ]
    return {
        "headline": [str(application["headline"]).strip()],
        "profile": [str(application["leadership_profile"]["text"]).strip()],
        "current_bullets": [str(item["text"]).strip() for item in application["current_role_bullets"]],
        "metric_values": [str(claim.get("value", "")) for claim in metrics],
        "metric_labels": [_metric_label(claim) for claim in metrics],
        "evidence_titles": [str(claim.get("label", "")) for claim in evidence_claims],
        "evidence_texts": [str(claim.get("safe_wording", "")) for claim in evidence_claims],
        "ksa_context": ["Representative KSA context: " + _recognized_context(packet)],
        "cube_bullets": cube_bullets,
        "earlier_bullets": earlier_bullets,
        "identity": [value for value in required_identity if value],
    }


def content_integrity_qa(expectations: dict[str, list[str]], docx_path: Path, pdf_path: Path, bundle: dict[str, Any]) -> dict[str, Any]:
    docx_text = _norm(_docx_text(docx_path))
    pdf_pages = _pdf_page_text(pdf_path)
    pdf_text = _norm("\n".join(pdf_pages))
    findings: list[dict[str, str]] = []

    def require(group: str, values: list[str], haystack: str, where: str) -> None:
        for value in values:
            if _norm(value) not in haystack:
                findings.append({"severity": "error", "code": "missing_in_" + where, "group": group, "value": value[:90]})

    for group in ("headline", "profile", "current_bullets", "metric_values", "metric_labels",
                  "evidence_titles", "evidence_texts", "ksa_context", "cube_bullets",
                  "earlier_bullets", "identity"):
        require(group, expectations[group], docx_text, "docx")
        require(group, expectations[group], pdf_text, "pdf")

    for anchor in FIXED_ANCHORS:
        if _norm(anchor) not in pdf_text:
            findings.append({"severity": "error", "code": "missing_fixed_anchor", "group": "anchors", "value": anchor})

    for placeholder in PLACEHOLDERS_ABSENT:
        if _norm(placeholder) in pdf_text or placeholder in _docx_text(docx_path):
            findings.append({"severity": "error", "code": "placeholder_present", "group": "placeholders", "value": placeholder})

    for dash in DASHES:
        if dash in "\n".join(pdf_pages) or dash in _docx_text(docx_path):
            findings.append({"severity": "error", "code": "forbidden_dash", "group": "characters", "value": repr(dash)})

    policy = bundle["config"]["policy"]
    for value in policy.get("prohibited_experience_names", []) + policy.get("prohibited_terms", []):
        if _norm(value) in pdf_text:
            findings.append({"severity": "error", "code": "prohibited_content", "group": "policy", "value": value})

    counts = {
        "metric_values": len(expectations["metric_values"]),
        "metric_labels": len(expectations["metric_labels"]),
        "evidence_cards": len(expectations["evidence_titles"]),
        "current_bullets": len(expectations["current_bullets"]),
        "cube_bullets": len(expectations["cube_bullets"]),
        "earlier_bullets": len(expectations["earlier_bullets"]),
    }
    if counts["metric_values"] != 6 or counts["metric_labels"] != 6:
        findings.append({"severity": "error", "code": "metric_count", "group": "counts", "value": json.dumps(counts)})
    if counts["evidence_cards"] != 4:
        findings.append({"severity": "error", "code": "evidence_count", "group": "counts", "value": json.dumps(counts)})
    return {"counts": counts, "findings": findings, "valid": not findings}


def _raster_page(pdf_path: Path, page: int, out_dir: Path, dpi: int = 100) -> Path:
    subprocess.run(
        ["pdftoppm", "-png", "-r", str(dpi), "-f", str(page), "-l", str(page), str(pdf_path), str(out_dir / f"p{page}")],
        capture_output=True, timeout=120, check=True,
    )
    matches = sorted(out_dir.glob(f"p{page}-*.png"))
    if not matches:
        raise RuntimeError(f"pdftoppm produced no image for page {page}")
    return matches[0]


def visual_qa(pdf_path: Path, template_reference_png: Path, *, root: Path | None = None) -> dict[str, Any]:
    from PIL import Image

    findings: list[dict[str, str]] = []
    info = subprocess.run(["pdfinfo", str(pdf_path)], capture_output=True, text=True, timeout=120, check=False)
    pages = 0
    width_pt = height_pt = 0.0
    for line in info.stdout.splitlines():
        if line.startswith("Pages:"):
            pages = int(line.split()[-1])
        if line.startswith("Page size:"):
            parts = line.split()
            width_pt, height_pt = float(parts[2]), float(parts[4])
    if pages != 2:
        findings.append({"severity": "error", "code": "page_count", "value": str(pages)})
    if abs(width_pt - 595.3) > 2 or abs(height_pt - 841.9) > 2:
        findings.append({"severity": "error", "code": "page_size_not_a4", "value": f"{width_pt}x{height_pt}"})

    with tempfile.TemporaryDirectory(prefix="sidebar-qa-", dir=_qa_temp_dir(root)) as tmp:
        tmp = Path(tmp)
        page1 = Image.open(_raster_page(pdf_path, 1, tmp)).convert("RGB")
        page2 = Image.open(_raster_page(pdf_path, 2, tmp)).convert("RGB")
        reference = Image.open(template_reference_png).convert("RGB")

        def ink_bounds(image: "Image.Image", threshold: int = 235) -> tuple[int, int, int, int]:
            gray = image.convert("L")
            pixels = gray.load()
            if pixels is None:  # pragma: no cover - PIL guarantees a loader
                raise RuntimeError("PIL image loader unavailable")
            width, height = gray.size
            min_x, min_y, max_x, max_y = width, height, -1, -1
            for y in range(0, height, 2):
                for x in range(0, width, 2):
                    if pixels[x, y] < threshold:
                        min_x, min_y = min(min_x, x), min(min_y, y)
                        max_x, max_y = max(max_x, x), max(max_y, y)
            return min_x, min_y, max_x, max_y

        # Page 2: all ink must stay inside the 1.0/1.1 cm margins (A4 @100dpi).
        margin_tb, margin_lr, tol = 39, 43, 14
        width2, height2 = page2.size
        b2 = ink_bounds(page2)
        if b2[0] < margin_lr - tol:
            findings.append({"severity": "error", "code": "page2_left_overflow", "value": str(b2[0])})
        if b2[1] < margin_tb - tol:
            findings.append({"severity": "error", "code": "page2_top_overflow", "value": str(b2[1])})
        if b2[2] > width2 - margin_lr + tol:
            findings.append({"severity": "error", "code": "page2_right_overflow", "value": str(b2[2])})
        if b2[3] > height2 - margin_tb + tol:
            findings.append({"severity": "error", "code": "page2_bottom_overflow", "value": str(b2[3])})
        if (b2[2] - b2[0]) < width2 // 3:
            findings.append({"severity": "error", "code": "page2_mostly_blank", "value": json.dumps(b2)})

        # Page 1: sidebar band present on the left; nothing right of the right margin.
        width1, height1 = page1.size
        b1 = ink_bounds(page1)
        if b1[0] > int(width1 * 0.08):
            findings.append({"severity": "error", "code": "page1_sidebar_missing", "value": str(b1[0])})
        if b1[2] > width1 - margin_lr + tol:
            findings.append({"severity": "error", "code": "page1_right_overflow", "value": str(b1[2])})

        # Headshot: crop the photo box from this render and from the template
        # reference; the approved headshot must still be there (color + match).
        band_right = _sidebar_band_right(page1)
        box = (int(band_right * 0.32), int(height1 * 0.035), int(band_right * 0.88), int(height1 * 0.165))
        crop = page1.crop(box)
        ref_crop = reference.crop(box)
        colorful = sum(1 for px in crop.getdata() if max(px) - min(px) > 30)
        if colorful < 300:
            findings.append({"severity": "error", "code": "headshot_missing", "value": f"colorful={colorful}"})
        diff = _mean_abs_diff(crop, ref_crop)
        if diff > 24.0:
            findings.append({"severity": "error", "code": "headshot_mismatch", "value": f"diff={diff:.1f}"})

    page_texts = _pdf_page_text(pdf_path)
    if len(page_texts) < 2 or "CONTINUED" not in page_texts[1].upper():
        findings.append({"severity": "error", "code": "page2_continuation_header_missing", "value": ""})
    return {
        "pages": pages,
        "size_pt": [width_pt, height_pt],
        "findings": findings,
        "valid": not findings,
    }


def _sidebar_band_right(image: "Image.Image", threshold: int = 120) -> int:
    gray = image.convert("L")
    pixels = gray.load()
    if pixels is None:  # pragma: no cover - PIL guarantees a loader
        raise RuntimeError("PIL image loader unavailable")
    width, height = gray.size
    band_right = 0
    for x in range(0, width // 2):
        dark = sum(1 for y in range(0, height, 6) if pixels[x, y] < threshold)
        if dark > (height // 6) * 0.7:
            band_right = x
    return max(band_right, 1)


def _mean_abs_diff(a: "Image.Image", b: "Image.Image") -> float:
    if a.size != b.size:
        b = b.resize(a.size)
    pa, pb = list(a.getdata()), list(b.getdata())
    total = sum(abs(x[c] - y[c]) for x, y in zip(pa, pb) for c in range(3))
    return total / (len(pa) * 3)


def _qa_temp_dir(root: Path | None = None) -> Path:
    path = qa_dir(root) / ".tmp"
    path.mkdir(parents=True, exist_ok=True)
    return path


def template_reference(root: Path | None = None) -> Path:
    """Render the configured template once and cache its page-1 raster for headshot comparison."""
    config, paths = load_config(root)
    cache = qa_dir(root) / "template-reference"
    cache.mkdir(parents=True, exist_ok=True)
    reference_png = cache / "template-page1.png"
    expected = file_sha256(paths.repo_root / config["template"]["repository_path"])
    stamp = cache / "template.sha256"
    if reference_png.is_file() and stamp.is_file() and stamp.read_text().strip() == expected:
        return reference_png
    with tempfile.TemporaryDirectory(prefix="template-ref-", dir=_qa_temp_dir(root)) as tmp:
        converted = convert_docx_to_pdf(paths.repo_root / config["template"]["repository_path"], Path(tmp))
        if not converted.get("converted"):
            raise RuntimeError("Template reference render failed: " + json.dumps(converted)[:300])
        page1 = _raster_page(Path(converted["pdf"]), 1, Path(tmp))
        shutil_copy(page1, reference_png)
    stamp.write_text(expected + "\n", encoding="utf-8")
    return reference_png


def shutil_copy(source: Path, destination: Path) -> None:
    import shutil

    shutil.copyfile(source, destination)


def enumerate_sidebar_jobs(*, root: Path | None = None) -> list[dict[str, Any]]:
    _, paths = load_config(root)
    tracker_module = _load_tracker_module(paths)
    tracker = tracker_module.CareerTracker(paths.tracker_base)
    jobs: list[dict[str, Any]] = []
    for row in tracker.list_rows():
        record = tracker.get_job(row["job_id"])
        artifacts = record.get("generated_artifacts") or []
        sidebar_pdfs = [
            item for item in artifacts
            if isinstance(item, dict) and item.get("variant") == "modern-executive-sidebar" and item.get("type") == "final_pdf"
        ]
        if not sidebar_pdfs:
            continue
        jobs.append({
            "job_id": row["job_id"],
            "company": row["company"],
            "role": row["role"],
            "processing_status": row.get("processing_status", ""),
            "application_status": row.get("application_status", ""),
            "existing_sidebar_revisions": len(sidebar_pdfs),
        })
    return jobs


def _load_tracker_module(paths: Any):
    import importlib.util

    module_path = paths.tracker_base / "tracker.py"
    spec = importlib.util.spec_from_file_location("live_career_tracker", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load tracker module from {module_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def revision_stem(outward_pdf: str, version: str) -> str:
    return Path(outward_pdf).stem + f"_v{version}"


def render_revision(job_id: str, *, root: Path | None = None) -> dict[str, Any]:
    config, paths = load_config(root)
    artifact_dir = paths.tracker_base / "artifacts" / job_id
    application = json.loads((artifact_dir / "generated_application.json").read_text(encoding="utf-8"))
    packet = json.loads((artifact_dir / "generation_packet.json").read_text(encoding="utf-8"))
    version = str(config["template"]["version"])
    revised = deepcopy(packet)
    revised["outward_filename"] = revision_stem(packet["outward_filename"], version) + ".docx"
    docx_info = render_docx(job_id, application, revised, root=root)
    docx_path = Path(docx_info["docx"])
    converted = convert_docx_to_pdf(docx_path, docx_path.parent)
    fit: dict[str, Any] = {"dense_split": {"applied": False}, "density_levels_applied": []}
    if converted.get("converted"):
        page_limit = int(str(config["template"]["page_limit"]))
        converted, fit = ensure_page_fit(docx_path, converted, page_limit=page_limit)
        docx_info["sha256"] = file_sha256(docx_path)
    if not converted.get("converted"):
        return {"job_id": job_id, "passed": False, "stage": "convert", "conversion": converted}
    return {
        "job_id": job_id,
        "passed": None,
        "stage": "qa",
        "docx": {"path": str(docx_path), "sha256": file_sha256(docx_path)},
        "pdf": {"path": converted["pdf"], "sha256": converted.get("sha256", file_sha256(Path(converted["pdf"])))},
        "application": application,
        "packet": packet,
        "fit": fit,
    }


def qa_job(job_id: str, *, root: Path | None = None, rendered: dict[str, Any] | None = None) -> dict[str, Any]:
    config, _ = load_config(root)
    bundle = load_bundle(root)
    rendered = rendered or render_revision(job_id, root=root)
    if rendered.get("stage") == "convert":
        report = {"job_id": job_id, "passed": False, "conversion": rendered.get("conversion"), "findings": [{"severity": "error", "code": "pdf_conversion_failed"}]}
        _write_qa_report(report, root=root)
        return report
    application = rendered["application"]
    packet = rendered["packet"]
    docx_path = Path(rendered["docx"]["path"])
    pdf_path = Path(rendered["pdf"]["path"])
    expectations = source_expectations(application, packet, root=root)
    content = content_integrity_qa(expectations, docx_path, pdf_path, bundle)
    pdf_check = verify_pdf(pdf_path, root=root)
    visual = visual_qa(pdf_path, template_reference(root), root=root)
    findings = list(content["findings"]) + list(visual["findings"])
    if not pdf_check.get("valid"):
        for item in pdf_check.get("findings", []):
            findings.append({"severity": "error", "code": "verify_pdf:" + item.get("code", "unknown"), "value": item.get("message", "")[:90]})
    report = {
        "job_id": job_id,
        "company": (application.get("job_id") or ""),
        "template_version": str(config["template"]["version"]),
        "template_sha256": config["template"]["expected_sha256"],
        "bundle_hash": str(packet.get("bundle_hash", "")),
        "docx": {"path": rendered["docx"]["path"], "sha256": rendered["docx"]["sha256"]},
        "pdf": {"path": rendered["pdf"]["path"], "sha256": rendered["pdf"]["sha256"], "page_count": pdf_check.get("page_count")},
        "content_integrity": content,
        "verify_pdf": {k: v for k, v in pdf_check.items() if k != "findings"},
        "visual": visual,
        "fit": rendered.get("fit", {}),
        "findings": findings,
        "passed": not findings,
        "checked_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
    }
    _write_qa_report(report, root=root)
    return report


def _write_qa_report(report: dict[str, Any], *, root: Path | None = None) -> None:
    target = qa_dir(root) / f"{report['job_id']}.json"
    target.parent.mkdir(parents=True, exist_ok=True)
    temp = target.with_suffix(".json.tmp")
    temp.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    temp.replace(target)


def apply_revision(job_id: str, *, root: Path | None = None, actor: str = "system") -> dict[str, Any]:
    config, paths = load_config(root)
    report_path = qa_dir(root) / f"{job_id}.json"
    if not report_path.is_file():
        return {"job_id": job_id, "applied": False, "blocker": "qa_report_missing"}
    report = json.loads(report_path.read_text(encoding="utf-8"))
    if not report.get("passed"):
        return {"job_id": job_id, "applied": False, "blocker": "qa_failed", "findings": report.get("findings", [])[:5]}
    docx_entry = {
        "type": "final_docx",
        "variant": "modern-executive-sidebar",
        "path": report["docx"]["path"],
        "sha256": report["docx"]["sha256"],
        "bundle_hash": report.get("bundle_hash", ""),
        "template_version": report.get("template_version", ""),
    }
    pdf_entry = {
        "type": "final_pdf",
        "variant": "modern-executive-sidebar",
        "path": report["pdf"]["path"],
        "sha256": report["pdf"]["sha256"],
        "bundle_hash": report.get("bundle_hash", ""),
        "template_version": report.get("template_version", ""),
    }
    tracker_module = _load_tracker_module(paths)
    tracker = tracker_module.CareerTracker(paths.tracker_base)
    record = tracker.get_job(job_id)
    existing = list(record.get("generated_artifacts") or [])
    if any(item.get("sha256") == pdf_entry["sha256"] and item.get("type") == "final_pdf" for item in existing):
        return {"job_id": job_id, "applied": True, "already_recorded": True}
    tracker.update_job(
        job_id,
        {"generated_artifacts": existing + [docx_entry, pdf_entry]},
        comment=(
            f"Regenerated Modern Executive Sidebar revision with approved template "
            f"v{report.get('template_version')} (sha {report.get('template_sha256', '')[:12]}); "
            "content-integrity and rendered-page QA passed; prior artifacts and submission evidence preserved unchanged"
        ),
        actor=actor,
        action="generated",
        source_refs=[docx_entry["path"], pdf_entry["path"]],
        confidence="high",
        requires_owner_review=False,
    )
    return {"job_id": job_id, "applied": True, "docx": docx_entry["path"], "pdf": pdf_entry["path"]}


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)
    sub.add_parser("plan")
    render = sub.add_parser("render")
    render.add_argument("--job-id", default="")
    render.add_argument("--all", action="store_true")
    render.add_argument("--limit", type=int, default=0)
    render.add_argument("--only", default="", help="comma-separated job ids")
    apply = sub.add_parser("apply")
    apply.add_argument("--job-id", default="")
    apply.add_argument("--all", action="store_true")
    args = parser.parse_args(argv)

    if args.command == "plan":
        print(json.dumps(enumerate_sidebar_jobs(), indent=2))
        return 0

    if args.command == "render":
        if args.all:
            jobs = enumerate_sidebar_jobs()
            if args.only:
                wanted = {item.strip() for item in args.only.split(",") if item.strip()}
                jobs = [item for item in jobs if item["job_id"] in wanted]
            if args.limit:
                jobs = jobs[: args.limit]
        else:
            if not args.job_id:
                parser.error("render requires --job-id or --all")
            jobs = [{"job_id": args.job_id}]
        failures = 0
        for item in jobs:
            report = qa_job(item["job_id"])
            status = "PASS" if report.get("passed") else "FAIL"
            if not report.get("passed"):
                failures += 1
            print(json.dumps({"job_id": item["job_id"], "qa": status, "findings": report.get("findings", [])[:6]}))
        return 1 if failures else 0

    if args.command == "apply":
        if args.all:
            jobs = [item["job_id"] for item in enumerate_sidebar_jobs()]
        else:
            if not args.job_id:
                parser.error("apply requires --job-id or --all")
            jobs = [args.job_id]
        failures = 0
        for job_id in jobs:
            result = apply_revision(job_id)
            if not result.get("applied"):
                failures += 1
            print(json.dumps(result))
        return 1 if failures else 0
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
